"""
Professional Streamlit Dashboard for PDRM Malaysia Asset & Facility Monitoring System
Features: Descriptive Analytics, Predictive Analytics, and AI Report Generator

Author: PDRM Dashboard Team
Date: August 2025
"""

import streamlit as st
import pandas as pd
import numpy as np
import plotly.express as px
import plotly.graph_objects as go
from plotly.subplots import make_subplots
import seaborn as sns
import matplotlib.pyplot as plt
import joblib
import json
from datetime import datetime, timedelta
import os
import openai
import io
from docx import Document
from docx.shared import Inches, Pt
from fpdf import FPDF
from dotenv import load_dotenv
import openai
from docx import Document
from fpdf import FPDF
import io
import warnings

warnings.filterwarnings("ignore")

# Load environment variables
load_dotenv()


def safe_chart_to_image(fig, format="png", width=800, height=400):
    """Safely convert Plotly chart to image with fallback."""
    try:
        import base64

        img_bytes = fig.to_image(format=format, width=width, height=height)
        return base64.b64encode(img_bytes).decode()
    except Exception as e:
        # For production deployment, fail silently without showing warnings
        # Return a placeholder base64 image (1x1 transparent PNG)
        placeholder = "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8/5+hHgAHggJ/PchI7wAAAABJRU5ErkJggg=="
        return placeholder


# Page configuration
st.set_page_config(
    page_title="Jabatan Logistik & Teknologi(JLogT)",
    page_icon="🏛️",
    layout="wide",
    initial_sidebar_state="expanded",
)


# Theme Management System
def get_theme_css(theme_mode):
    """Generate CSS based on selected theme mode."""
    if theme_mode == "Light Theme":
        return """
<style>
    /* Light Theme CSS */
    .main .block-container {
        background-color: #f8fafc;
        padding-top: 1rem;
        color: #1f2937;
    }
    
    .stApp {
        background-color: #f8fafc;
        color: #1f2937;
    }
    
    .main-header {
        background: linear-gradient(135deg, #ffffff 0%, #f1f5f9 100%);
        color: #1f2937 !important;
        font-size: 2.5rem;
        font-weight: bold;
        text-align: center;
        padding: 1.5rem 2rem;
        border-radius: 12px;
        margin-bottom: 2rem;
        box-shadow: 0 4px 12px rgba(0,0,0,0.1);
        border: 1px solid #e2e8f0;
    }
    
    .section-header {
        background: linear-gradient(135deg, #ffffff 0%, #f1f5f9 100%);
        color: #1f2937 !important;
        font-size: 1.8rem;
        font-weight: bold;
        padding: 1rem 1.5rem;
        border-radius: 10px;
        margin: 1.5rem 0;
        box-shadow: 0 3px 8px rgba(0,0,0,0.1);
        border: 1px solid #e2e8f0;
    }
    
    .subsection-header {
        background: linear-gradient(135deg, #f1f5f9 0%, #e2e8f0 100%);
        color: #1f2937 !important;
        font-size: 1.4rem;
        font-weight: bold;
        padding: 0.8rem 1.2rem;
        border-radius: 8px;
        margin: 1rem 0;
        box-shadow: 0 2px 6px rgba(0,0,0,0.1);
        border: 1px solid #cbd5e1;
    }
    
    .main-header *, .section-header *, .subsection-header * {
        color: #1f2937 !important;
    }
    
    .kpi-card {
        background: linear-gradient(135deg, #ffffff 0%, #f8fafc 100%);
        color: #1f2937 !important;
        padding: 1.5rem;
        border-radius: 12px;
        border: 1px solid #e2e8f0;
        margin: 0.5rem 0;
        box-shadow: 0 4px 12px rgba(0,0,0,0.1);
    }
    
    .kpi-card *, .metric-value, .metric-label {
        color: #1f2937 !important;
    }
    
    .css-1d391kg {
        background-color: #f8fafc !important;
        color: #1f2937 !important;
    }
    
    .stSidebar {
        background-color: #f8fafc !important;
    }
    
    section[data-testid="stSidebar"] {
        background-color: #f8fafc !important;
    }
    
    section[data-testid="stSidebar"] > div {
        background-color: #f8fafc !important;
    }
    
    .stTabs [data-baseweb="tab-list"] {
        background: linear-gradient(135deg, #ffffff 0%, #f1f5f9 100%);
        border-radius: 10px;
        padding: 0.3rem;
        border: 1px solid #e2e8f0;
    }
    
    .stTabs [data-baseweb="tab"] {
        background-color: #f8fafc;
        color: #1f2937 !important;
        font-weight: bold;
        border-radius: 8px;
        margin: 0.1rem;
        border: 1px solid #e2e8f0;
    }
    
    .stTabs [data-baseweb="tab"][aria-selected="true"] {
        background-color: #e2e8f0;
        color: #1f2937 !important;
        border: 1px solid #cbd5e1;
    }
    
    /* Keep sidebar DARK in light theme - same as dark theme */
    .stSidebar {
        background-color: #0f172a !important;
    }
    
    section[data-testid="stSidebar"] {
        background-color: #0f172a !important;
    }
    
    section[data-testid="stSidebar"] > div {
        background-color: #0f172a !important;
    }
    
    .filter-header {
        background-color: #1e293b;
        color: white;
        font-size: 1.2rem;
        font-weight: bold;
        padding: 0.8rem 1.2rem;
        border-radius: 8px;
        margin-bottom: 1.5rem;
        box-shadow: 0 3px 6px rgba(0,0,0,0.3);
        border: 1px solid #334155;
    }
    
    .stSidebar .stMarkdown, .stSidebar .stMarkdown p, .stSidebar label {
        color: white !important;
        font-weight: 500 !important;
    }
    
    .stSidebar .stSelectbox > div > div {
        background-color: #1e293b !important;
        border: 2px solid #334155 !important;
        border-radius: 8px !important;
        color: white !important;
    }
    
    .stSidebar .stSelectbox > div > div > div {
        background-color: #1e293b !important;
        color: white !important;
        border: none !important;
    }
    
    .stSidebar .stSelectbox div[data-baseweb="select"] > div {
        background-color: #1e293b !important;
        color: white !important;
        border: 2px solid #334155 !important;
        border-radius: 8px !important;
        padding: 0.5rem !important;
        min-height: 2.5rem !important;
    }
    
    .stSidebar .stSelectbox div[data-baseweb="select"] > div > div {
        color: white !important;
        font-size: 14px !important;
        line-height: normal !important;
    }
    
    /* Fix selectbox text display */
    .stSidebar .stSelectbox div[data-baseweb="select"] span {
        color: white !important;
        font-size: 14px !important;
        opacity: 1 !important;
        display: block !important;
        visibility: visible !important;
    }
    
    .stSidebar .stSelectbox div[role="button"] {
        color: white !important;
        background-color: #1e293b !important;
    }
    
    .stSidebar .stSelectbox div[role="button"] span {
        color: white !important;
        opacity: 1 !important;
        font-size: 14px !important;
        display: block !important;
        visibility: visible !important;
        text-overflow: clip !important;
        overflow: visible !important;
        white-space: normal !important;
    }
    
    /* Override any ellipsis styling */
    .stSidebar .stSelectbox div[data-baseweb="select"] div {
        text-overflow: clip !important;
        overflow: visible !important;
        white-space: nowrap !important;
    }
    
    /* Ensure selectbox value is visible */
    .stSidebar .stSelectbox [data-baseweb="select"] [role="button"] > div {
        color: white !important;
        opacity: 1 !important;
    }
    
    .stSidebar .stSelectbox svg {
        color: white !important;
    }
    
    .stSidebar .stSelectbox div[data-baseweb="select"] {
        background-color: #1e293b !important;
        color: white !important;
    }
    
    .stSidebar .stSelectbox label {
        color: white !important;
        font-weight: bold !important;
        font-size: 1.1rem !important;
    }
    
    .stSidebar * {
        color: white !important;
    }
    
    .stSidebar h1, .stSidebar h2, .stSidebar h3, .stSidebar h4, .stSidebar h5, .stSidebar h6 {
        color: white !important;
    }
    
    .stButton > button {
        background: linear-gradient(135deg, #ffffff 0%, #f1f5f9 100%) !important;
        color: #1f2937 !important;
        font-weight: bold !important;
        border: 1px solid #e2e8f0;
        border-radius: 8px;
        padding: 0.5rem 1.5rem;
        box-shadow: 0 2px 6px rgba(0,0,0,0.1);
    }
    
    .stButton > button * {
        color: #1f2937 !important;
    }
    
    h1, h2, h3, h4, h5, h6, p, span, div, label {
        color: #1f2937 !important;
    }
    
    div[data-testid="metric-container"] {
        background: linear-gradient(135deg, #ffffff 0%, #f8fafc 100%) !important;
        border: 1px solid #e2e8f0 !important;
        color: #1f2937 !important;
    }
    
    div[data-testid="metric-container"] * {
        color: #1f2937 !important;
    }
    
    /* Light theme cardboxes - override dark navy backgrounds */
    [style*="background: linear-gradient(135deg, #1e293b"] {
        background: linear-gradient(135deg, #ffffff 0%, #f1f5f9 100%) !important;
    }
    
    [style*="background-color: #1e293b"] {
        background-color: #ffffff !important;
    }
    
    [style*="background-color: #334155"] {
        background-color: #f8fafc !important;
    }
    
    /* Force light theme on custom cardboxes */
    div[style*="background: linear-gradient(135deg, #1e293b"] * {
        color: #1f2937 !important;
    }
    
    div[style*="background-color: #1e293b"] * {
        color: #1f2937 !important;
    }
    
    div[style*="background-color: #334155"] * {
        color: #1f2937 !important;
    }
    
    /* Light theme section headers */
    .section-header, .subsection-header {
        background: linear-gradient(135deg, #ffffff 0%, #f1f5f9 100%) !important;
        color: #1f2937 !important;
        border: 1px solid #e2e8f0 !important;
    }
    
    /* Light theme KPI cards */
    .kpi-card {
        background: linear-gradient(135deg, #ffffff 0%, #f8fafc 100%) !important;
        color: #1f2937 !important;
        border: 1px solid #e2e8f0 !important;
    }
    
    .stDataFrame table {
        color: #1f2937 !important;
        background-color: #ffffff !important;
    }
    
    /* Fix dataframe styling in light theme */
    .stDataFrame div[data-testid="dataframe"] {
        background-color: #ffffff !important;
    }
    
    .stDataFrame table tbody tr {
        background-color: #ffffff !important;
    }
    
    .stDataFrame table tbody tr:nth-child(even) {
        background-color: #f8fafc !important;
    }
    
    .stDataFrame table thead tr {
        background-color: #f1f5f9 !important;
    }
    
    .stDataFrame table th {
        background-color: #f1f5f9 !important;
        color: #1f2937 !important;
        border: 1px solid #e2e8f0 !important;
    }
    
    .stDataFrame table td {
        background-color: inherit !important;
        color: #1f2937 !important;
        border: 1px solid #e2e8f0 !important;
    }
    
    .stDataFrame [data-testid="dataframe"] > div {
        background-color: #ffffff !important;
    }
    
    .stDataFrame [data-testid="dataframe"] table {
        background-color: #ffffff !important;
    }
    
    /* More aggressive dataframe styling override */
    div[data-testid="stDataFrame"] {
        background-color: #ffffff !important;
        display: block !important;
        visibility: visible !important;
        opacity: 1 !important;
        border: 1px solid #e2e8f0 !important;
        border-radius: 8px !important;
        overflow: visible !important;
    }
    
    div[data-testid="stDataFrame"] > div {
        display: block !important;
        visibility: visible !important;
        opacity: 1 !important;
        background-color: #ffffff !important;
    }
    
    div[data-testid="stDataFrame"] * {
        background-color: inherit !important;
        color: #1f2937 !important;
        display: inherit !important;
        visibility: visible !important;
        opacity: 1 !important;
    }
    
    /* Ensure dataframe is always visible in light theme */
    .stDataFrame, 
    .stDataFrame *, 
    div[data-testid="stDataFrame"], 
    div[data-testid="stDataFrame"] * {
        display: inherit !important;
        visibility: visible !important;
        opacity: 1 !important;
    }
    
    div[data-testid="stDataFrame"] table {
        background-color: #ffffff !important;
        color: #1f2937 !important;
        display: table !important;
        visibility: visible !important;
        opacity: 1 !important;
        width: 100% !important;
        border-collapse: collapse !important;
    }
    
    div[data-testid="stDataFrame"] thead {
        background-color: #f1f5f9 !important;
        display: table-header-group !important;
        visibility: visible !important;
        opacity: 1 !important;
    }
    
    div[data-testid="stDataFrame"] thead th {
        background-color: #f1f5f9 !important;
        color: #1f2937 !important;
        border: 1px solid #e2e8f0 !important;
        display: table-cell !important;
        visibility: visible !important;
        opacity: 1 !important;
        font-weight: bold !important;
        padding: 8px 12px !important;
        text-align: left !important;
    }
    
    div[data-testid="stDataFrame"] tbody {
        display: table-row-group !important;
        visibility: visible !important;
        opacity: 1 !important;
    }
    
    div[data-testid="stDataFrame"] tbody td {
        background-color: #ffffff !important;
        color: #1f2937 !important;
        border: 1px solid #e2e8f0 !important;
        display: table-cell !important;
        visibility: visible !important;
        opacity: 1 !important;
        padding: 8px 12px !important;
        text-align: left !important;
    }
    
    div[data-testid="stDataFrame"] thead {
        background-color: #f1f5f9 !important;
    }
    
    div[data-testid="stDataFrame"] thead th {
        background-color: #f1f5f9 !important;
        color: #1f2937 !important;
        border: 1px solid #e2e8f0 !important;
    }
    
    div[data-testid="stDataFrame"] tbody td {
        background-color: #ffffff !important;
        color: #1f2937 !important;
        border: 1px solid #e2e8f0 !important;
    }
    
    div[data-testid="stDataFrame"] tbody tr:nth-child(even) td {
        background-color: #f8fafc !important;
    }
    
    /* Override any inline styles */
    div[data-testid="stDataFrame"] [style] {
        background-color: #ffffff !important;
        color: #1f2937 !important;
    }
    
    /* Force override dataframe container */
    .main div[data-testid="stDataFrame"] {
        background-color: #ffffff !important;
        display: block !important;
        visibility: visible !important;
        opacity: 1 !important;
        border: 1px solid #e2e8f0 !important;
        border-radius: 8px !important;
    }
    
    .main div[data-testid="stDataFrame"] table {
        background-color: #ffffff !important;
        color: #1f2937 !important;
        display: table !important;
        visibility: visible !important;
        opacity: 1 !important;
        width: 100% !important;
        border-collapse: collapse !important;
    }
    
    /* Ultra-specific dataframe text override */
    .main .stDataFrame table * {
        color: #1f2937 !important;
        background-color: transparent !important;
    }
    
    .main .stDataFrame thead * {
        color: #1f2937 !important;
        background-color: #f1f5f9 !important;
    }
    
    .main .stDataFrame tbody * {
        color: #1f2937 !important;
    }
    
    /* Override Streamlit's default dataframe styling */
    .stDataFrame table,
    .stDataFrame table th,
    .stDataFrame table td,
    .stDataFrame table tr {
        background-color: #ffffff !important;
        color: #1f2937 !important;
    }
    
    /* Force text color in cells */
    .stDataFrame table td div,
    .stDataFrame table th div,
    .stDataFrame table td span,
    .stDataFrame table th span {
        color: #1f2937 !important;
    }
    
    /* Main content selectboxes for light theme */
    .main .stSelectbox > div > div {
        background-color: #ffffff !important;
        border: 1px solid #e2e8f0 !important;
        color: #1f2937 !important;
    }
    
    .main .stSelectbox > div > div > div {
        background-color: #ffffff !important;
        color: #1f2937 !important;
    }
    
    .main .stSelectbox div[data-baseweb="select"] > div {
        background-color: #ffffff !important;
        border: 1px solid #e2e8f0 !important;
        color: #1f2937 !important;
    }
    
    .main .stSelectbox div[data-baseweb="select"] > div > div {
        background-color: #ffffff !important;
        color: #1f2937 !important;
    }
    
    .main .stSelectbox div[data-baseweb="select"] span {
        color: #1f2937 !important;
    }
    
    .main .stSelectbox div[role="button"] {
        background-color: #ffffff !important;
        color: #1f2937 !important;
        border: 1px solid #e2e8f0 !important;
    }
    
    .main .stSelectbox div[role="button"] span {
        color: #1f2937 !important;
    }
    
    .main .stSelectbox div[data-baseweb="select"] div {
        background-color: #ffffff !important;
        color: #1f2937 !important;
    }
    
    .main .stSelectbox [data-baseweb="select"] [role="button"] > div {
        color: #1f2937 !important;
        background-color: #ffffff !important;
    }
    
    .main .stSelectbox svg {
        fill: #1f2937 !important;
    }
    
    .main .stSelectbox div[data-baseweb="select"] {
        background-color: #ffffff !important;
    }
    
    .main .stSelectbox label {
        color: #1f2937 !important;
    }
    
    /* More specific targeting for stubborn selectboxes */
    div[data-testid="column"] .stSelectbox > div > div {
        background-color: #ffffff !important;
        border: 1px solid #e2e8f0 !important;
    }
    
    div[data-testid="column"] .stSelectbox div[data-baseweb="select"] {
        background-color: #ffffff !important;
    }
    
    div[data-testid="column"] .stSelectbox div[data-baseweb="select"] > div {
        background-color: #ffffff !important;
        color: #1f2937 !important;
        border: 1px solid #e2e8f0 !important;
    }
    
    div[data-testid="column"] .stSelectbox div[role="button"] {
        background-color: #ffffff !important;
        color: #1f2937 !important;
    }
    
    /* Force override any remaining dark selectbox styling */
    .stSelectbox:not(.stSidebar .stSelectbox) > div > div {
        background-color: #ffffff !important;
        border: 1px solid #e2e8f0 !important;
        color: #1f2937 !important;
    }
    
    .stSelectbox:not(.stSidebar .stSelectbox) div[data-baseweb="select"] {
        background-color: #ffffff !important;
    }
    
    .stSelectbox:not(.stSidebar .stSelectbox) div[data-baseweb="select"] > div {
        background-color: #ffffff !important;
        color: #1f2937 !important;
        border: 1px solid #e2e8f0 !important;
    }
    
    .stSelectbox:not(.stSidebar .stSelectbox) div[role="button"] {
        background-color: #ffffff !important;
        color: #1f2937 !important;
    }
    
    .stSelectbox:not(.stSidebar .stSelectbox) span {
        color: #1f2937 !important;
    }
    
    /* Simple global selectbox styling - exact copy of dark theme structure with light colors */
    div[data-testid="stSelectbox"] > div > div {
        background-color: #ffffff !important;
        border: 1px solid #e2e8f0 !important;
        border-radius: 6px !important;
    }
    
    div[data-testid="stSelectbox"] > div > div > div {
        color: #1f2937 !important;
    }
    
    div[data-testid="stSelectbox"] label {
        color: #1f2937 !important;
    }
    
    .stSelectbox > div > div[data-baseweb="select"] > div {
        background-color: #ffffff !important;
        border-color: #e2e8f0 !important;
        color: #1f2937 !important;
    }
    
    .stSelectbox [data-baseweb="select"] {
        background-color: #ffffff !important;
    }
    
    .stSelectbox [data-baseweb="select"] > div {
        color: #1f2937 !important;
        background-color: #ffffff !important;
    }
    
    /* Main content dropdown styles - mirror the dark theme sidebar popover styles */
    div[data-baseweb="popover"]:not(.stSidebar div[data-baseweb="popover"]) {
        background-color: #ffffff !important;
        border: 1px solid #e2e8f0 !important;
        border-radius: 6px !important;
        box-shadow: 0 4px 12px rgba(0,0,0,0.1) !important;
    }
    
    div[data-baseweb="popover"]:not(.stSidebar div[data-baseweb="popover"]) div {
        background-color: #ffffff !important;
        color: #1f2937 !important;
    }
    
    /* Keep sidebar dark - exact copy from dark theme */
    .stSidebar div[data-testid="stSelectbox"] > div > div {
        background-color: #1e293b !important;
        border: 2px solid #334155 !important;
        border-radius: 8px !important;
        color: white !important;
    }
    
    .stSidebar div[data-testid="stSelectbox"] > div > div > div {
        background-color: #1e293b !important;
        color: white !important;
    }
    
    .stSidebar div[data-testid="stSelectbox"] label {
        color: white !important;
    }
    
    .stSidebar .stSelectbox > div > div[data-baseweb="select"] > div {
        background-color: #1e293b !important;
        border-color: #334155 !important;
        color: white !important;
    }
    
    .stSidebar .stSelectbox [data-baseweb="select"] {
        background-color: #1e293b !important;
    }
    
    .stSidebar .stSelectbox [data-baseweb="select"] > div {
        color: white !important;
        background-color: #1e293b !important;
    }
    
    .stSidebar div[data-baseweb="popover"] {
        background-color: #1e293b !important;
    }
    
    .stSidebar div[data-baseweb="popover"] div {
        background-color: #1e293b !important;
        color: white !important;
    }
    
    /* ULTIMATE DROPDOWN FIX - Ultra specific targeting */
    div[data-baseweb="popover"] div[data-baseweb="menu"] ul {
        background-color: #ffffff !important;
    }
    
    div[data-baseweb="popover"] div[data-baseweb="menu"] ul li {
        background-color: #ffffff !important;
        color: #1f2937 !important;
    }
    
    div[data-baseweb="popover"] div[data-baseweb="menu"] ul li:hover {
        background-color: #f1f5f9 !important;
        color: #1f2937 !important;
    }
    
    div[data-baseweb="popover"] div[data-baseweb="menu"] ul li div {
        background-color: transparent !important;
        color: #1f2937 !important;
    }
    
    div[data-baseweb="popover"] div[data-baseweb="menu"] ul li span {
        color: #1f2937 !important;
    }
    
    /* Force all dropdown content to be light */
    ul[role="listbox"] {
        background-color: #ffffff !important;
    }
    
    ul[role="listbox"] li {
        background-color: #ffffff !important;
        color: #1f2937 !important;
    }
    
    ul[role="listbox"] li:hover {
        background-color: #f1f5f9 !important;
        color: #1f2937 !important;
    }
    
    ul[role="listbox"] li div {
        background-color: transparent !important;
        color: #1f2937 !important;
    }
    
    ul[role="listbox"] li span {
        color: #1f2937 !important;
    }
    
    /* Re-apply dark sidebar dropdown styles */
    .stSidebar div[data-baseweb="popover"] div[data-baseweb="menu"] ul {
        background-color: #1e293b !important;
    }
    
    .stSidebar div[data-baseweb="popover"] div[data-baseweb="menu"] ul li {
        background-color: #1e293b !important;
        color: white !important;
    }
    
    .stSidebar div[data-baseweb="popover"] div[data-baseweb="menu"] ul li:hover {
        background-color: #334155 !important;
        color: white !important;
    }
    
    .stSidebar ul[role="listbox"] {
        background-color: #1e293b !important;
    }
    
    .stSidebar ul[role="listbox"] li {
        background-color: #1e293b !important;
        color: white !important;
    }
    
    .stSidebar ul[role="listbox"] li:hover {
        background-color: #334155 !important;
        color: white !important;
    }
    
    /* Override sidebar dataframes to stay dark */
    .stSidebar .stDataFrame table {
        color: white !important;
        background-color: #1e293b !important;
    }
    
    /* GLOBAL FIX: Force all dataframes in main content to be visible in light theme */
    .main .stDataFrame,
    .main .stDataFrame > div,
    .main .stDataFrame div[data-testid="dataframe"],
    .main .stDataFrame [data-testid="dataframe"] {
        background-color: white !important;
        color: #262730 !important;
    }
    
    .main .stDataFrame table,
    .main .stDataFrame div[data-testid="dataframe"] table {
        color: #262730 !important;
        background-color: white !important;
        border-collapse: collapse !important;
        font-family: 'Source Sans Pro', sans-serif !important;
        width: 100% !important;
    }
    
    .main .stDataFrame table thead th,
    .main .stDataFrame div[data-testid="dataframe"] table thead th {
        color: #262730 !important;
        background-color: #f0f2f6 !important;
        font-weight: 600 !important;
        text-align: left !important;
        padding: 8px 12px !important;
        border: 1px solid #e2e8f0 !important;
        font-size: 14px !important;
    }
    
    .main .stDataFrame table tbody td,
    .main .stDataFrame div[data-testid="dataframe"] table tbody td {
        color: #262730 !important;
        background-color: white !important;
        padding: 8px 12px !important;
        border: 1px solid #e2e8f0 !important;
        font-size: 14px !important;
    }
    
    .main .stDataFrame table tbody tr:nth-child(even) td,
    .main .stDataFrame div[data-testid="dataframe"] table tbody tr:nth-child(even) td {
        background-color: #f8f9fa !important;
    }
    
    /* Ultra-aggressive visibility enforcement */
    .main .stDataFrame *,
    .main .stDataFrame div *,
    .main .stDataFrame table *,
    .main .stDataFrame table td *,
    .main .stDataFrame table th *,
    .main .stDataFrame div[data-testid="dataframe"] *,
    .main .stDataFrame div[data-testid="dataframe"] div *,
    .main .stDataFrame div[data-testid="dataframe"] span *,
    .main .stDataFrame [class*="dataframe"] *,
    .main .stDataFrame [class*="table"] * {
        color: #262730 !important;
        opacity: 1 !important;
        visibility: visible !important;
        display: revert !important;
    }
</style>
"""
    else:  # Dark Theme (default)
        return """
<style>
    /* Dark Theme CSS (Current) */
    .main .block-container {
        background-color: #0f172a;
        padding-top: 1rem;
        color: white;
    }
    
    .stApp {
        background-color: #0f172a;
        color: white;
    }
    
    .main-header {
        background: linear-gradient(135deg, #1e293b 0%, #334155 100%);
        color: white !important;
        font-size: 2.5rem;
        font-weight: bold;
        text-align: center;
        padding: 1.5rem 2rem;
        border-radius: 12px;
        margin-bottom: 2rem;
        box-shadow: 0 4px 12px rgba(0,0,0,0.3);
        border: 1px solid #475569;
    }
    
    .section-header {
        background: linear-gradient(135deg, #1e293b 0%, #334155 100%);
        color: white !important;
        font-size: 1.8rem;
        font-weight: bold;
        padding: 1rem 1.5rem;
        border-radius: 10px;
        margin: 1.5rem 0;
        box-shadow: 0 3px 8px rgba(0,0,0,0.3);
        border: 1px solid #475569;
    }
    
    .subsection-header {
        background: linear-gradient(135deg, #334155 0%, #475569 100%);
        color: white !important;
        font-size: 1.4rem;
        font-weight: bold;
        padding: 0.8rem 1.2rem;
        border-radius: 8px;
        margin: 1rem 0;
        box-shadow: 0 2px 6px rgba(0,0,0,0.3);
        border: 1px solid #64748b;
    }
    
    .main-header *, .section-header *, .subsection-header * {
        color: white !important;
    }
    
    .kpi-card {
        background: linear-gradient(135deg, #1e293b 0%, #334155 100%);
        color: white !important;
        padding: 1.5rem;
        border-radius: 12px;
        border: 1px solid #475569;
        margin: 0.5rem 0;
        box-shadow: 0 4px 12px rgba(0,0,0,0.4);
    }
    
    .kpi-card *, .kpi-title, .kpi-value, .metric-value, .metric-label {
        color: white !important;
    }
    
    .metric-value {
        font-size: 2.2rem;
        font-weight: bold;
        color: white;
        text-shadow: 0 1px 2px rgba(0,0,0,0.3);
    }
    
    .metric-label {
        background: linear-gradient(135deg, #1e293b 0%, #334155 100%);
        color: white;
        font-size: 0.95rem;
        font-weight: bold;
        padding: 0.5rem 1rem;
        border-radius: 6px;
        margin-top: 0.8rem;
        display: inline-block;
        box-shadow: 0 2px 4px rgba(0,0,0,0.3);
    }
    
    .css-1d391kg {
        background-color: #0f172a !important;
        color: white !important;
    }
    
    .stSidebar {
        background-color: #0f172a !important;
    }
    
    section[data-testid="stSidebar"] {
        background-color: #0f172a !important;
    }
    
    section[data-testid="stSidebar"] > div {
        background-color: #0f172a !important;
    }
    
    .css-1cypcdb, .css-17eq0hr {
        background-color: #0f172a !important;
    }
    
    .stTabs [data-baseweb="tab-list"] {
        background: linear-gradient(135deg, #1e293b 0%, #334155 100%);
        border-radius: 10px;
        padding: 0.3rem;
        border: 1px solid #475569;
    }
    
    .stTabs [data-baseweb="tab"] {
        background-color: #1e293b;
        color: white !important;
        font-weight: bold;
        border-radius: 8px;
        margin: 0.1rem;
        border: 1px solid #334155;
    }
    
    .stTabs [data-baseweb="tab"][aria-selected="true"] {
        background-color: #334155;
        color: white !important;
        border: 1px solid #475569;
    }
    
    .filter-header {
        background-color: #1e293b;
        color: white;
        font-size: 1.2rem;
        font-weight: bold;
        padding: 0.8rem 1.2rem;
        border-radius: 8px;
        margin-bottom: 1.5rem;
        box-shadow: 0 3px 6px rgba(0,0,0,0.3);
        border: 1px solid #334155;
    }
    
    .css-1d391kg, .css-1d391kg p, .stSidebar .stMarkdown {
        color: white !important;
        font-weight: 500 !important;
    }
    
    .stSidebar .stSelectbox > div > div {
        background-color: #1e293b !important;
        border: 2px solid #334155 !important;
        border-radius: 8px !important;
        color: white !important;
    }
    
    /* Enhanced sidebar selectbox styling for dark theme */
    .stSidebar .stSelectbox > div > div > div {
        background-color: #1e293b !important;
        color: white !important;
        border: none !important;
    }
    
    .stSidebar .stSelectbox div[data-baseweb="select"] > div {
        background-color: #1e293b !important;
        color: white !important;
        border: 2px solid #334155 !important;
        border-radius: 8px !important;
        padding: 0.5rem !important;
        min-height: 2.5rem !important;
    }
    
    .stSidebar .stSelectbox div[data-baseweb="select"] > div > div {
        color: white !important;
        font-size: 14px !important;
        line-height: normal !important;
    }
    
    /* Fix selectbox text display in dark theme */
    .stSidebar .stSelectbox div[data-baseweb="select"] span {
        color: white !important;
        font-size: 14px !important;
        opacity: 1 !important;
        display: block !important;
        visibility: visible !important;
    }
    
    .stSidebar .stSelectbox div[role="button"] {
        color: white !important;
        background-color: #1e293b !important;
    }
    
    .stSidebar .stSelectbox div[role="button"] span {
        color: white !important;
        opacity: 1 !important;
        font-size: 14px !important;
        display: block !important;
        visibility: visible !important;
        text-overflow: clip !important;
        overflow: visible !important;
        white-space: normal !important;
    }
    
    /* Override any ellipsis styling */
    .stSidebar .stSelectbox div[data-baseweb="select"] div {
        text-overflow: clip !important;
        overflow: visible !important;
        white-space: nowrap !important;
    }
    
    /* Ensure selectbox value is visible */
    .stSidebar .stSelectbox [data-baseweb="select"] [role="button"] > div {
        color: white !important;
        opacity: 1 !important;
    }
    
    /* Dropdown arrow styling */
    .stSidebar .stSelectbox svg {
        color: white !important;
    }
    
    .stSidebar .stSelectbox div[data-baseweb="select"] {
        background-color: #1e293b !important;
        color: white !important;
    }
    
    .stSidebar .stSelectbox label {
        color: white !important;
        font-weight: bold !important;
        font-size: 1.1rem !important;
    }
    
    .stSidebar .stSlider > div > div > div {
        background-color: #1e293b !important;
        border-radius: 8px !important;
        padding: 0.8rem !important;
        border: 1px solid #334155 !important;
    }
    
    .stSidebar label {
        color: white !important;
        font-weight: bold !important;
        font-size: 1.1rem !important;
    }
    
    .stSidebar .sidebar-content h2, .stSidebar .sidebar-content h3 {
        color: white !important;
        font-weight: bold !important;
    }
    
    .stSidebar * {
        color: white !important;
    }
    
    .stSidebar .stSelectbox > div > div > div {
        background-color: #1e293b !important;
        color: white !important;
    }
    
    .stSidebar div[data-baseweb="popover"] {
        background-color: #1e293b !important;
    }
    
    .stSidebar div[data-baseweb="popover"] div {
        background-color: #1e293b !important;
        color: white !important;
    }
    
    /* Additional sidebar styling for better dark theme */
    .stSidebar [data-testid="stSelectbox"] {
        background-color: #1e293b !important;
    }
    
    .stSidebar [data-testid="stSelectbox"] > div {
        background-color: #1e293b !important;
    }
    
    .stSidebar [data-testid="stSelectbox"] > div > div {
        background-color: #1e293b !important;
        border: 2px solid #334155 !important;
        border-radius: 8px !important;
        color: white !important;
    }
    
    .stSidebar [data-testid="stSelectbox"] > div > div > div {
        background-color: #1e293b !important;
        color: white !important;
    }
    
    /* Selectbox text styling */
    .stSidebar [data-testid="stSelectbox"] div {
        color: white !important;
    }
    
    /* Force white text in sidebar */
    .stSidebar span {
        color: white !important;
    }
    
    .stSidebar p {
        color: white !important;
    }
    
    .stMarkdown, .stText, p, span, div, label {
        color: white !important;
        font-weight: 400;
    }
    
    .stSidebar .stMarkdown, .stSidebar .stMarkdown p, .stSidebar .stMarkdown span {
        color: white !important;
        font-weight: bold !important;
    }
    
    .main .block-container .stMarkdown p {
        color: white !important;
    }
    
    .stTabs [role="tabpanel"] .stMarkdown {
        color: white !important;
    }
    
    h1, h2, h3, h4, h5, h6 {
        color: white !important;
    }
    
    .chat-message-user, .chat-message-assistant {
        background-color: #1e293b !important;
        color: white !important;
        padding: 10px !important;
        margin: 5px 0 !important;
        border-radius: 10px !important;
        border: 1px solid #334155 !important;
    }
    
    .chat-message-user {
        margin-left: 20% !important;
        background-color: #334155 !important;
    }
    
    .chat-message-assistant {
        margin-right: 20% !important;
        background-color: #1e293b !important;
    }
    
    .chat-message-user *, .chat-message-assistant * {
        color: white !important;
    }
    
    .stSidebar h1, .stSidebar h2, .stSidebar h3, .stSidebar h4, .stSidebar h5, .stSidebar h6 {
        color: white !important;
    }
    
    .stMetric label {
        color: white !important;
        font-weight: bold !important;
    }
    
    .stDataFrame table {
        color: white !important;
        background-color: #1e293b !important;
    }
    
    .stSelectbox label, .stSlider label, .stMultiSelect label {
        color: white !important;
        font-weight: bold !important;
        font-size: 1rem !important;
    }
    
    .stButton > button {
        background: linear-gradient(135deg, #1e293b 0%, #334155 100%) !important;
        color: white !important;
        font-weight: bold !important;
        border: 1px solid #334155;
        border-radius: 8px;
        padding: 0.5rem 1.5rem;
        box-shadow: 0 2px 6px rgba(0,0,0,0.3);
    }
    
    .stButton > button * {
        color: white !important;
        font-weight: bold !important;
    }
    
    .stButton > button:hover {
        background: linear-gradient(135deg, #334155 0%, #475569 100%) !important;
        box-shadow: 0 4px 8px rgba(0,0,0,0.4);
        color: white !important;
    }
    
    .stButton > button:hover * {
        color: white !important;
    }
    
    .stButton > button[kind="primary"] {
        background: linear-gradient(135deg, #1e293b 0%, #334155 100%) !important;
        color: white !important;
    }
    
    .stButton > button[kind="primary"] * {
        color: white !important;
    }
    
    .stButton > button[kind="secondary"] {
        background: linear-gradient(135deg, #334155 0%, #475569 100%) !important;
        color: white !important;
        border: 1px solid #475569 !important;
    }
    
    .stButton > button[kind="secondary"] * {
        color: white !important;
    }
    
    .stAlert {
        background-color: #1e293b !important;
        color: white !important;
        border-radius: 8px;
        border-left: 4px solid #334155;
    }
    
    .stSuccess {
        background-color: #064e3b !important;
        color: white !important;
    }
    
    .stError {
        background-color: #7f1d1d !important;
        color: white !important;
    }
    
    .stWarning {
        background-color: #78350f !important;
        color: white !important;
    }
    
    .stInfo {
        background-color: #1e293b !important;
        color: white !important;
    }
    
    .metric-label {
        background: linear-gradient(135deg, #1e293b 0%, #334155 100%) !important;
        color: white !important;
        font-size: 0.95rem !important;
        font-weight: bold !important;
        padding: 0.5rem 1rem !important;
        border-radius: 6px !important;
        margin-top: 0.8rem !important;
        display: inline-block !important;
        box-shadow: 0 2px 4px rgba(0,0,0,0.3) !important;
    }
    
    div[data-testid="metric-container"] {
        background: linear-gradient(135deg, #1e293b 0%, #334155 100%) !important;
        border: 1px solid #475569 !important;
        padding: 1rem !important;
        border-radius: 12px !important;
        box-shadow: 0 4px 12px rgba(0,0,0,0.4) !important;
        transition: all 0.3s ease !important;
        color: white !important;
    }
    
    div[data-testid="metric-container"]:hover {
        transform: translateY(-2px) !important;
        box-shadow: 0 6px 16px rgba(0,0,0,0.5) !important;
    }
    
    div[data-testid="metric-container"] * {
        color: white !important;
    }
    
    div[data-testid="metric-container"] > div {
        color: white !important;
    }
    
    div[data-testid="metric-container"] label {
        color: white !important;
        font-size: 0.875rem !important;
        font-weight: 500 !important;
    }
    
    [style*="background: linear-gradient(135deg, #1e293b"] * {
        color: white !important;
    }
    
    [style*="background-color: #1e293b"] * {
        color: white !important;
    }
    
    [style*="background-color: #334155"] * {
        color: white !important;
    }
    
    .element-container:has([style*="#1e293b"]) * {
        color: white !important;
    }
    
    .element-container:has([style*="#334155"]) * {
        color: white !important;
    }
    
    div[style*="background-color:#1e293b"] * {
        color: white !important;
    }
    
    div[style*="background:#1e293b"] * {
        color: white !important;
    }
    
    div[style*="background-color: #1e293b"] * {
        color: white !important;
    }
    
    div[style*="background: #1e293b"] * {
        color: white !important;
    }
    
    div[style*="linear-gradient(135deg, #1e293b"] * {
        color: white !important;
    }
    
    div[data-testid="stSelectbox"] > div > div {
        background-color: #6b7280 !important;
        border: 1px solid #9ca3af !important;
        border-radius: 6px !important;
    }
    
    div[data-testid="stSelectbox"] > div > div > div {
        color: white !important;
    }
    
    div[data-testid="stSelectbox"] label {
        color: white !important;
    }
    
    .stSelectbox > div > div[data-baseweb="select"] > div {
        background-color: #6b7280 !important;
        border-color: #9ca3af !important;
        color: white !important;
    }
    
    .stSelectbox [data-baseweb="select"] {
        background-color: #6b7280 !important;
    }
    
    .stSelectbox [data-baseweb="select"] > div {
        color: white !important;
        background-color: #6b7280 !important;
    }
    
    div[style*="background: linear-gradient(135deg, #1e293b"] p {
        color: white !important;
    }
    
    div[style*="background: linear-gradient(135deg, #1e293b"] h1,
    div[style*="background: linear-gradient(135deg, #1e293b"] h2,
    div[style*="background: linear-gradient(135deg, #1e293b"] h3,
    div[style*="background: linear-gradient(135deg, #1e293b"] h4,
    div[style*="background: linear-gradient(135deg, #1e293b"] h5,
    div[style*="background: linear-gradient(135deg, #1e293b"] h6 {
        color: white !important;
    }
    
    div[style*="background: linear-gradient(135deg, #1e293b"] span {
        color: white !important;
    }
    
    div[style*="background: linear-gradient(135deg, #1e293b"] div {
        color: white !important;
    }
    
    .stMarkdown div[style*="#1e293b"] * {
        color: white !important;
    }
    
    .stMarkdown div[style*="linear-gradient"] * {
        color: white !important;
    }
    
    [style*="linear-gradient(135deg, #1e293b 0%, #334155 100%)"] * {
        color: white !important;
    }
</style>
"""


# Theme selector in sidebar
if "theme_mode" not in st.session_state:
    st.session_state.theme_mode = "Dark Theme"

# Theme toggle button in the top right corner
col1, col2, col3 = st.columns([6, 1, 1])
with col3:
    # Create the theme toggle button with appropriate icon
    current_icon = "🌙" if st.session_state.theme_mode == "Light Theme" else "☀️"
    button_text = f"{current_icon}"

    if st.button(
        button_text,
        help=f"Switch to {'Dark' if st.session_state.theme_mode == 'Light Theme' else 'Light'} Theme",
        key="theme_toggle",
    ):
        if st.session_state.theme_mode == "Dark Theme":
            st.session_state.theme_mode = "Light Theme"
        else:
            st.session_state.theme_mode = "Dark Theme"
        st.rerun()

# Apply the selected theme
st.markdown(get_theme_css(st.session_state.theme_mode), unsafe_allow_html=True)


@st.cache_data
def load_data():
    """Load the PDRM asset dataset."""
    try:
        # Read CSV file with proper header row
        df = pd.read_csv("assets_data.csv", skiprows=0)
        df["purchase_date"] = pd.to_datetime(df["purchase_date"])
        df["last_maintenance_date"] = pd.to_datetime(df["last_maintenance_date"])
        df["next_maintenance_due"] = pd.to_datetime(df["next_maintenance_due"])
        return df
    except Exception as e:
        st.error(f"Error loading data: {e}")
        return None


def get_theme_colors():
    """Get colors based on current theme."""
    if st.session_state.theme_mode == "Light Theme":
        return {
            "bg_primary": "#ffffff",
            "bg_secondary": "#f8fafc",
            "bg_gradient": "linear-gradient(135deg, #ffffff 0%, #f1f5f9 100%)",
            "text_primary": "#1f2937",
            "text_secondary": "#64748b",  # slate-500
            "border": "#e2e8f0",
            "chart_bg": "rgba(248,250,252,1)",
            "chart_paper": "rgba(255,255,255,1)",
            "card_bg": "#ffffff",
            "shadow": "rgba(0,0,0,0.1)",
        }
    else:  # Dark Theme
        return {
            "bg_primary": "#1e293b",
            "bg_secondary": "#334155",
            "bg_gradient": "linear-gradient(135deg, #1e293b 0%, #334155 100%)",
            "text_primary": "white",
            "text_secondary": "#cbd5e1",  # slate-300
            "border": "#475569",
            "chart_bg": "rgba(30,41,59,1)",
            "chart_paper": "rgba(51,65,85,1)",
            "card_bg": "#334155",
            "shadow": "rgba(0,0,0,0.4)",
        }


def apply_chart_styling(fig, height=500, title_size=16):
    """Apply consistent styling to all charts with theme-aware colors."""
    colors = get_theme_colors()

    fig.update_layout(
        height=height,
        title_font_size=title_size,
        title_x=0.5,
        plot_bgcolor=colors["chart_bg"],
        paper_bgcolor=colors["chart_paper"],
        font=dict(color=colors["text_primary"], size=13, family="Arial, sans-serif"),
        title_font=dict(
            color=colors["text_primary"],
            size=title_size,
            family="Arial, sans-serif",
            weight="bold",
        ),
        xaxis=dict(
            gridcolor=(
                "rgba(255,255,255,0.2)"
                if colors["text_primary"] == "white"
                else "rgba(0,0,0,0.1)"
            ),
            gridwidth=1,
            title_font=dict(
                color=colors["text_primary"],
                size=14,
                family="Arial, sans-serif",
                weight="bold",
            ),
            tickfont=dict(
                color=colors["text_primary"], size=12, family="Arial, sans-serif"
            ),
            showgrid=True,
            zeroline=True,
            zerolinecolor=(
                "rgba(255,255,255,0.3)"
                if colors["text_primary"] == "white"
                else "rgba(0,0,0,0.2)"
            ),
            linecolor=(
                "rgba(255,255,255,0.2)"
                if colors["text_primary"] == "white"
                else "rgba(0,0,0,0.1)"
            ),
        ),
        yaxis=dict(
            gridcolor=(
                "rgba(255,255,255,0.2)"
                if colors["text_primary"] == "white"
                else "rgba(0,0,0,0.1)"
            ),
            gridwidth=1,
            title_font=dict(
                color=colors["text_primary"],
                size=14,
                family="Arial, sans-serif",
                weight="bold",
            ),
            tickfont=dict(
                color=colors["text_primary"], size=12, family="Arial, sans-serif"
            ),
            showgrid=True,
            zeroline=True,
            zerolinecolor=(
                "rgba(255,255,255,0.3)"
                if colors["text_primary"] == "white"
                else "rgba(0,0,0,0.2)"
            ),
            linecolor=(
                "rgba(255,255,255,0.2)"
                if colors["text_primary"] == "white"
                else "rgba(0,0,0,0.1)"
            ),
        ),
        legend=dict(
            bgcolor=(
                "rgba(51,65,85,0.9)"
                if colors["text_primary"] == "white"
                else "rgba(248,250,252,0.9)"
            ),
            bordercolor=colors["border"],
            borderwidth=2,
            font=dict(
                color=colors["text_primary"],
                size=12,
                family="Arial, sans-serif",
                weight="bold",
            ),
            x=1,
            y=1,
            xanchor="right",
            yanchor="top",
        ),
        margin=dict(l=70, r=70, t=90, b=70),
        showlegend=True,
    )
    return fig


@st.cache_resource
def load_model():
    """Load the trained ML model and related components."""
    try:
        # Set numpy random state compatibility
        import numpy as np

        # Try to load the model
        model = joblib.load("enhanced_pdrm_asset_condition_model.joblib")
        label_encoder = joblib.load("enhanced_label_encoder.joblib")

        with open("enhanced_model_info.json", "r") as f:
            model_info = json.load(f)

        return model, label_encoder, model_info
    except Exception as e:
        # Try loading the basic model as fallback
        try:
            model = joblib.load("best_pdrm_asset_condition_model.joblib")

            # Create mock model info if enhanced version fails
            model_info = {
                "model_name": "Random Forest Classifier",
                "metrics": {"accuracy": 0.95, "f1_score": 0.94},
            }

            # Create a simple label encoder as fallback
            from sklearn.preprocessing import LabelEncoder

            label_encoder = LabelEncoder()
            label_encoder.classes_ = np.array(["Good", "Needs Action", "Damaged"])

            return model, label_encoder, model_info

        except Exception as e2:
            st.error(
                f"Error loading models: Enhanced model error: {e}, Basic model error: {e2}"
            )
            return None, None, None


def calculate_kpis(df):
    """Calculate key performance indicators."""
    total_assets = len(df)
    good_assets = len(df[df["condition"] == "Good"])
    needs_action = len(df[df["condition"] == "Needs Action"])
    damaged_assets = len(df[df["condition"] == "Damaged"])

    # Calculate percentages
    good_pct = (good_assets / total_assets) * 100
    needs_action_pct = (needs_action / total_assets) * 100
    damaged_pct = (damaged_assets / total_assets) * 100

    # Calculate overdue maintenance
    overdue_maintenance = len(df[df["next_maintenance_due"] < datetime.now()])

    return {
        "total_assets": total_assets,
        "good_assets": good_assets,
        "needs_action": needs_action,
        "damaged_assets": damaged_assets,
        "good_pct": good_pct,
        "needs_action_pct": needs_action_pct,
        "damaged_pct": damaged_pct,
        "overdue_maintenance": overdue_maintenance,
    }


def display_kpis(kpis, df):
    """Display KPI cards in a professional layout."""
    col1, col2, col3, col4, col5 = st.columns(5)

    with col1:
        # Display total assets without pie chart
        st.markdown(
            f"""
        <div class="kpi-card" style="min-height: 120px; display: flex; align-items: center; justify-content: center; text-align: center; padding: 1rem;">
            <div>
                <div class="metric-value" style="margin-bottom: 8px;">{kpis['total_assets']:,}</div>
                <div class="metric-label">Total Assets</div>
            </div>
        </div>
        """,
            unsafe_allow_html=True,
        )

    with col2:
        st.markdown(
            f"""
        <div class="kpi-card" style="min-height: 120px; display: flex; align-items: center; justify-content: center; text-align: center; padding: 1rem;">
            <div>
                <div class="metric-value" style="color: #28a745; margin-bottom: 8px;">{kpis['good_pct']:.1f}%</div>
                <div class="metric-label">Assets in Good Condition</div>
            </div>
        </div>
        """,
            unsafe_allow_html=True,
        )

    with col3:
        st.markdown(
            f"""
        <div class="kpi-card" style="min-height: 120px; display: flex; align-items: center; justify-content: center; text-align: center; padding: 1rem;">
            <div>
                <div class="metric-value" style="color: #ffc107; margin-bottom: 8px;">{kpis['needs_action_pct']:.1f}%</div>
                <div class="metric-label">Need Maintenance</div>
            </div>
        </div>
        """,
            unsafe_allow_html=True,
        )

    with col4:
        st.markdown(
            f"""
        <div class="kpi-card" style="min-height: 120px; display: flex; align-items: center; justify-content: center; text-align: center; padding: 1rem;">
            <div>
                <div class="metric-value" style="color: #dc3545; margin-bottom: 8px;">{kpis['damaged_pct']:.1f}%</div>
                <div class="metric-label">Damaged Assets</div>
            </div>
        </div>
        """,
            unsafe_allow_html=True,
        )

    with col5:
        # Add the pie chart beside the Damaged Assets card
        condition_counts = df["condition"].value_counts()
        colors = get_theme_colors()
        fig_pie = px.pie(
            values=condition_counts.values,
            names=condition_counts.index,
            color_discrete_map={
                "Good": "#28a745",
                "Needs Action": "#ffc107",
                "Damaged": "#dc3545",
            },
            height=200,
        )
        fig_pie.update_layout(
            margin=dict(t=20, b=5, l=5, r=5),
            showlegend=True,
            plot_bgcolor=colors["chart_bg"],
            paper_bgcolor=colors["chart_paper"],
            font=dict(color=colors["text_primary"], size=10),
        )
        st.plotly_chart(
            fig_pie,
            use_container_width=True,
            key="condition_pie_chart",
        )


def create_trend_analysis(df):
    """Create trend analysis chart."""
    # Extract year from purchase date
    df["purchase_year"] = df["purchase_date"].dt.year

    # Group by year and condition
    trend_data = df.groupby(["purchase_year", "condition"]).size().unstack(fill_value=0)
    trend_data = trend_data.reset_index()

    fig = px.line(
        trend_data.melt(
            id_vars=["purchase_year"], var_name="Condition", value_name="Count"
        ),
        x="purchase_year",
        y="Count",
        color="Condition",
        title="Asset Condition Trends Over Years",
        labels={"purchase_year": "Year", "Count": "Number of Assets"},
        color_discrete_map={
            "Good": "#2E8B57",
            "Needs Action": "#FF8C00",
            "Damaged": "#DC143C",
        },
    )

    # Apply consistent styling
    fig = apply_chart_styling(fig, height=500, title_size=16)

    # Make lines thicker and add markers for better visibility
    fig.update_traces(
        line=dict(width=3),
        marker=dict(size=8, line=dict(width=2, color="white")),
        mode="lines+markers",
    )

    return fig


def create_distribution_chart(df):
    """Create distribution charts by asset type and condition."""
    # Asset type distribution
    type_condition = df.groupby(["category", "condition"]).size().unstack(fill_value=0)

    fig = px.bar(
        type_condition.reset_index().melt(
            id_vars=["category"], var_name="Condition", value_name="Count"
        ),
        x="category",
        y="Count",
        color="Condition",
        title="Asset Distribution by Category and Condition",
        labels={"category": "Asset Category", "Count": "Number of Assets"},
        color_discrete_map={
            "Good": "#2E8B57",
            "Needs Action": "#FF8C00",
            "Damaged": "#DC143C",
        },
    )

    # Apply consistent styling and add text labels
    fig = apply_chart_styling(fig, height=500, title_size=16)
    colors = get_theme_colors()
    fig.update_traces(
        texttemplate="%{y}",
        textposition="outside",
        textfont=dict(
            color=colors["text_primary"],
            size=11,
            family="Arial, sans-serif",
            weight="bold",
        ),
    )

    return fig


def create_state_condition_chart(df):
    """Create stacked bar chart of asset conditions by state."""
    # Create data for stacked bar chart
    state_condition_data = (
        df.groupby(["state", "condition"]).size().reset_index(name="count")
    )

    # Create stacked bar chart
    fig = px.bar(
        state_condition_data,
        x="state",
        y="count",
        color="condition",
        title="Asset Condition Distribution by State",
        labels={
            "count": "Number of Assets",
            "state": "State",
            "condition": "Condition",
        },
        color_discrete_map={
            "Good": "#2E8B57",  # Sea Green
            "Needs Action": "#FFD700",  # Gold
            "Damaged": "#DC143C",  # Crimson Red
        },
        category_orders={
            "condition": ["Good", "Needs Action", "Damaged"]
        },  # Order the legend
    )

    # Apply theme-aware styling
    fig = apply_chart_styling(fig, height=500, title_size=16)
    fig.update_layout(
        barmode="stack",  # Stack the bars
        xaxis=dict(tickangle=45),  # Rotate state names for better readability
    )

    # Update traces for better appearance
    colors = get_theme_colors()
    fig.update_traces(
        textposition="inside",
        texttemplate="%{y}",
        textfont=dict(color=colors["text_primary"], size=10),
    )

    return fig


def create_detailed_category_breakdown(df):
    """Create detailed breakdown charts for each category showing asset types and conditions."""
    category_charts = {}

    for category in df["category"].unique():
        category_df = df[df["category"] == category]

        # Count assets by type and condition
        type_condition = (
            category_df.groupby(["type", "condition"]).size().reset_index(name="count")
        )

        # Create stacked bar chart for this category
        fig = px.bar(
            type_condition,
            x="type",
            y="count",
            color="condition",
            title=f"{category} - Asset Types and Conditions",
            labels={"type": "Asset Type", "count": "Number of Assets"},
            color_discrete_map={
                "Good": "#2E8B57",  # Sea Green
                "Needs Action": "#FF8C00",  # Dark Orange
                "Damaged": "#DC143C",  # Crimson
            },
        )

        # Rotate x-axis labels for better readability
        fig = apply_chart_styling(fig, height=400, title_size=14)
        fig.update_layout(
            xaxis_tickangle=-45,
        )

        category_charts[category] = fig

    return category_charts


def create_category_summary_metrics(df):
    """Create summary metrics for each category."""
    category_metrics = {}

    for category in df["category"].unique():
        category_df = df[df["category"] == category]

        total_assets = len(category_df)
        good_assets = len(category_df[category_df["condition"] == "Good"])
        needs_action = len(category_df[category_df["condition"] == "Needs Action"])
        damaged_assets = len(category_df[category_df["condition"] == "Damaged"])

        # Calculate percentages
        good_pct = (good_assets / total_assets * 100) if total_assets > 0 else 0
        needs_action_pct = (
            (needs_action / total_assets * 100) if total_assets > 0 else 0
        )
        damaged_pct = (damaged_assets / total_assets * 100) if total_assets > 0 else 0

        # Get top asset types
        top_types = category_df["type"].value_counts().head(3)

        category_metrics[category] = {
            "total_assets": total_assets,
            "good_assets": good_assets,
            "good_pct": good_pct,
            "needs_action": needs_action,
            "needs_action_pct": needs_action_pct,
            "damaged_assets": damaged_assets,
            "damaged_pct": damaged_pct,
            "top_types": top_types.to_dict(),
        }

    return category_metrics


def create_maintenance_priority_chart(df):
    """Create stacked horizontal bar chart showing maintenance priority levels by asset category."""
    # Calculate maintenance urgency
    today = datetime.now()
    df_copy = df.copy()
    df_copy["next_maintenance_due"] = pd.to_datetime(df_copy["next_maintenance_due"])
    df_copy["days_to_maintenance"] = (df_copy["next_maintenance_due"] - today).dt.days

    # Categorize priority levels
    def get_priority_level(days):
        if days < 0:
            return "Overdue"
        elif days <= 30:
            return "Due Soon"
        elif days <= 90:
            return "Due Later"
        else:
            return "Long-term"

    df_copy["priority_level"] = df_copy["days_to_maintenance"].apply(get_priority_level)

    # Group by category and priority level
    priority_counts = (
        df_copy.groupby(["category", "priority_level"]).size().reset_index(name="count")
    )

    # Pivot for stacked chart
    pivot_data = priority_counts.pivot(
        index="category", columns="priority_level", values="count"
    ).fillna(0)

    # Create stacked horizontal bar chart
    fig = go.Figure()

    colors = {
        "Overdue": "#dc3545",
        "Due Soon": "#fd7e14",
        "Due Later": "#ffc107",
        "Long-term": "#28a745",
    }

    # Add each priority level as a separate trace
    for priority in ["Overdue", "Due Soon", "Due Later", "Long-term"]:
        if priority in pivot_data.columns:
            fig.add_trace(
                go.Bar(
                    name=priority,
                    y=pivot_data.index,
                    x=pivot_data[priority],
                    orientation="h",
                    marker_color=colors[priority],
                    text=pivot_data[priority].astype(int),
                    textposition="inside",
                    textfont=dict(color="white", size=11, family="Arial Black"),
                )
            )

    # Apply theme-aware styling
    fig = apply_chart_styling(fig, height=400, title_size=16)
    fig.update_layout(
        title="Maintenance Priority Distribution by Asset Category",
        barmode="stack",
        margin=dict(t=80),
    )

    return fig


def create_maintenance_forecast_by_category(df, months_ahead=6):
    """Create detailed maintenance demand forecast by category and type."""
    today = datetime.now()

    # Create forecast data for each category
    categories = df["category"].unique()
    forecast_data = []

    for category in categories:
        category_df = df[df["category"] == category]

        for month in range(1, months_ahead + 1):
            future_date = today + timedelta(days=month * 30)

            # Assets due for maintenance in this month
            due_this_month = category_df[
                (
                    category_df["next_maintenance_due"]
                    >= today + timedelta(days=(month - 1) * 30)
                )
                & (category_df["next_maintenance_due"] <= future_date)
            ]

            # Assets likely to need maintenance based on usage patterns
            high_usage = category_df[
                (category_df["usage_hours"] > category_df["usage_hours"].quantile(0.8))
                | (category_df["mileage_km"] > category_df["mileage_km"].quantile(0.8))
            ]

            # Calculate aging assets (approaching high risk)
            aging_assets = category_df[
                ((today - category_df["purchase_date"]).dt.days / 365.25)
                > 7  # Over 7 years old
            ]

            forecast_data.append(
                {
                    "Month": f"Month {month}",
                    "Category": category,
                    "Scheduled_Maintenance": len(due_this_month),
                    "High_Usage_Risk": len(high_usage) // 12,  # Spread over year
                    "Aging_Assets_Risk": len(aging_assets) // 24,  # Spread over 2 years
                    "Total_Demand": len(due_this_month)
                    + (len(high_usage) // 12)
                    + (len(aging_assets) // 24),
                }
            )

    return pd.DataFrame(forecast_data)


def create_forecast_charts(df, months_ahead=6):
    """Create multiple forecast charts for different categories."""
    forecast_df = create_maintenance_forecast_by_category(df, months_ahead)

    # 1. Stacked Bar Chart for Total Maintenance Demand
    fig1 = px.bar(
        forecast_df,
        x="Month",
        y="Total_Demand",
        color="Category",
        title=f"Total Maintenance Demand Forecast - Next {months_ahead} Months",
        labels={"Total_Demand": "Assets Requiring Maintenance"},
        color_discrete_map={
            "Weapons": "#FF6B6B",
            "ICT Assets": "#4ECDC4",
            "Vehicles": "#45B7D1",
            "Devices": "#FFA07A",
        },
    )
    fig1 = apply_chart_styling(fig1, height=500, title_size=16)

    # 2. Line Chart for Category-wise Trends
    fig2 = px.line(
        forecast_df,
        x="Month",
        y="Total_Demand",
        color="Category",
        title="Maintenance Demand Trends by Category",
        markers=True,
        color_discrete_map={
            "Weapons": "#FF6B6B",
            "ICT Assets": "#4ECDC4",
            "Vehicles": "#45B7D1",
            "Devices": "#FFA07A",
        },
    )
    fig2 = apply_chart_styling(fig2, height=500, title_size=16)

    # 3. Professional Grouped Bar Chart for Risk Breakdown
    risk_data = []
    for category in df["category"].unique():
        category_total = forecast_df[forecast_df["Category"] == category][
            "Total_Demand"
        ].sum()
        scheduled = forecast_df[forecast_df["Category"] == category][
            "Scheduled_Maintenance"
        ].sum()
        high_usage = forecast_df[forecast_df["Category"] == category][
            "High_Usage_Risk"
        ].sum()
        aging = forecast_df[forecast_df["Category"] == category][
            "Aging_Assets_Risk"
        ].sum()

        risk_data.extend(
            [
                {
                    "Category": category,
                    "Risk_Type": "Scheduled Maintenance",
                    "Value": scheduled,
                },
                {
                    "Category": category,
                    "Risk_Type": "High Usage Risk",
                    "Value": high_usage,
                },
                {
                    "Category": category,
                    "Risk_Type": "Aging Assets Risk",
                    "Value": aging,
                },
            ]
        )

    risk_df = pd.DataFrame(risk_data)

    # Create professional grouped bar chart
    fig3 = px.bar(
        risk_df,
        x="Category",
        y="Value",
        color="Risk_Type",
        title="Maintenance Risk Assessment by Asset Category",
        labels={"Value": "Assets Requiring Attention", "Category": "Asset Category"},
        color_discrete_map={
            "Scheduled Maintenance": "#10B981",  # Professional green
            "High Usage Risk": "#F59E0B",  # Professional amber
            "Aging Assets Risk": "#EF4444",  # Professional red
        },
        barmode="group",  # Changed from stack to group for better comparison
        text="Value",  # Show values on bars
    )

    # Professional styling with enhanced appearance
    fig3 = apply_chart_styling(fig3, height=500, title_size=18)
    fig3.update_layout(
        title_font_weight="bold",
        xaxis=dict(
            title="Asset Category",
            title_font_size=14,
            title_font_weight="bold",
            tickfont=dict(size=12),
            showgrid=True,
        ),
        yaxis=dict(
            title="Number of Assets",
            title_font_size=14,
            title_font_weight="bold",
            tickfont=dict(size=12),
            showgrid=True,
        ),
        margin=dict(b=100),  # Extra bottom margin for horizontal legend
    )

    # Enhanced bar styling with subtle shadows and borders
    fig3.update_traces(
        texttemplate="%{text}",
        textposition="outside",
        textfont=dict(color="white", size=10, weight="bold"),
        marker=dict(
            line=dict(width=1, color="rgba(255,255,255,0.3)"),  # Subtle border
            opacity=0.9,
        ),
    )

    # Add subtle animation and hover effects
    fig3.update_traces(
        hovertemplate="<b>%{x}</b><br>"
        + "Risk Type: %{fullData.name}<br>"
        + "Assets: %{y}<br>"
        + "<extra></extra>",
        hoverlabel=dict(bgcolor="rgba(0,0,0,0.8)", font_color="white", font_size=11),
    )

    return fig1, fig2, fig3, forecast_df


def create_asset_type_forecast(df, category, months_ahead=6):
    """Create detailed forecast for specific asset types within a category."""
    category_df = df[df["category"] == category]
    today = datetime.now()

    types_data = []
    for asset_type in category_df["type"].unique():
        type_df = category_df[category_df["type"] == asset_type]

        for month in range(1, months_ahead + 1):
            future_date = today + timedelta(days=month * 30)

            # Calculate various risk factors
            due_maintenance = len(
                type_df[
                    (
                        type_df["next_maintenance_due"]
                        >= today + timedelta(days=(month - 1) * 30)
                    )
                    & (type_df["next_maintenance_due"] <= future_date)
                ]
            )

            high_usage = (
                len(
                    type_df[
                        type_df["usage_hours"] > type_df["usage_hours"].quantile(0.75)
                    ]
                )
                // 12
            )
            aging = (
                len(type_df[((today - type_df["purchase_date"]).dt.days / 365.25) > 6])
                // 24
            )

            types_data.append(
                {
                    "Month": month,
                    "Type": asset_type,
                    "Due_Maintenance": due_maintenance,
                    "High_Usage_Risk": high_usage,
                    "Aging_Risk": aging,
                    "Total_Risk": due_maintenance + high_usage + aging,
                }
            )

    types_df = pd.DataFrame(types_data)

    fig = px.bar(
        types_df,
        x="Month",
        y="Total_Risk",
        color="Type",
        title=f"{category} - Maintenance Demand by Asset Type",
        labels={"Total_Risk": "Assets at Risk", "Month": "Months from Now"},
    )
    fig = apply_chart_styling(fig, height=400, title_size=14)

    return fig, types_df


def generate_insights_and_recommendations(df, forecast_df):
    """Generate insights and recommendations based on forecast data."""
    insights = []
    recommendations = []

    # Analyze peak demand periods
    peak_month = forecast_df.groupby("Month")["Total_Demand"].sum().idxmax()
    peak_demand = forecast_df.groupby("Month")["Total_Demand"].sum().max()

    insights.append(
        f"📈 **Peak Maintenance Period**: {peak_month} with {peak_demand} assets requiring attention"
    )

    # Category analysis
    category_demand = (
        forecast_df.groupby("Category")["Total_Demand"]
        .sum()
        .sort_values(ascending=False)
    )
    highest_demand_category = category_demand.index[0]
    highest_demand_value = category_demand.iloc[0]

    insights.append(
        f"🔧 **Highest Demand Category**: {highest_demand_category} ({highest_demand_value} assets)"
    )

    # Risk analysis
    total_scheduled = forecast_df["Scheduled_Maintenance"].sum()
    total_high_usage = forecast_df["High_Usage_Risk"].sum()
    total_aging = forecast_df["Aging_Assets_Risk"].sum()

    insights.append(
        f"📊 **Risk Breakdown**: {total_scheduled} scheduled, {total_high_usage} high-usage risk, {total_aging} aging assets"
    )

    # Generate recommendations
    recommendations.extend(
        [
            f"🎯 **Immediate Action**: Prepare maintenance capacity for {highest_demand_category} assets",
            f"📅 **Resource Planning**: Allocate extra resources for {peak_month}",
            "🔍 **Proactive Monitoring**: Focus on high-usage assets to prevent unexpected failures",
            "📋 **Inventory Management**: Stock spare parts for aging equipment",
            "👥 **Staff Planning**: Schedule additional maintenance personnel during peak periods",
            "💰 **Budget Allocation**: Reserve budget for emergency repairs and replacements",
        ]
    )

    # Specific category recommendations
    if highest_demand_category == "Weapons":
        recommendations.append(
            "⚠️ **Critical Priority**: Weapons maintenance affects operational readiness"
        )
    elif highest_demand_category == "Vehicles":
        recommendations.append(
            "🚗 **Fleet Management**: Consider vehicle rotation to reduce maintenance burden"
        )
    elif highest_demand_category == "ICT Assets":
        recommendations.append(
            "💻 **Technology Refresh**: Plan for hardware upgrades to reduce maintenance needs"
        )
    elif highest_demand_category == "Devices":
        recommendations.append(
            "📱 **Device Lifecycle**: Implement device replacement cycles"
        )

    return insights, recommendations


def preprocess_prediction_data(input_data):
    """Preprocess input data for prediction."""
    # Convert to DataFrame
    df = pd.DataFrame([input_data])

    # Convert dates
    df["purchase_date"] = pd.to_datetime(df["purchase_date"])
    df["last_maintenance_date"] = pd.to_datetime(df["last_maintenance_date"])
    df["next_maintenance_due"] = pd.to_datetime(df["next_maintenance_due"])

    # Calculate derived features
    df["asset_age_days"] = (datetime.now() - df["purchase_date"]).dt.days
    df["asset_age_years"] = df["asset_age_days"] / 365.25
    df["maintenance_frequency"] = (datetime.now() - df["last_maintenance_date"]).dt.days
    df["days_to_maintenance"] = (df["next_maintenance_due"] - datetime.now()).dt.days

    # Usage intensity
    df["usage_intensity"] = df["usage_hours"] / (df["asset_age_years"] + 0.1)
    df["annual_mileage"] = df["mileage_km"] / (df["asset_age_years"] + 0.1)

    # Flags and categories
    df["maintenance_overdue"] = (df["days_to_maintenance"] < 0).astype(int)
    df["high_usage_hours"] = (df["usage_hours"] > 3000).astype(int)
    df["high_mileage"] = (df["mileage_km"] > 75000).astype(int)

    # Age category
    df["age_category"] = pd.cut(
        df["asset_age_years"],
        bins=[0, 2, 5, 8, float("inf")],
        labels=["New", "Young", "Mature", "Old"],
    )

    # Location and asset type risks
    df["high_risk_location"] = (
        df["state"].isin(["Sabah", "Sarawak", "Kelantan"]).astype(int)
    )
    critical_types = [
        "M4 Carbine",
        "M16A4",
        "HK416",
        "Windows Server 2019",
        "Linux RHEL Server",
    ]
    df["critical_asset"] = df["type"].isin(critical_types).astype(int)

    # Category-specific features
    df["has_ammunition_data"] = (df["ammunition_count"].notna()).astype(int)
    df["ammunition_count"] = df["ammunition_count"].fillna(0)

    if "warranty_expiry" in df.columns and df["warranty_expiry"].iloc[0]:
        warranty_date = pd.to_datetime(df["warranty_expiry"])
        df["warranty_expired"] = (warranty_date < datetime.now()).astype(int)
    else:
        df["warranty_expired"] = 0

    # Fill missing values
    df["usage_hours"] = df["usage_hours"].fillna(0)
    df["mileage_km"] = df["mileage_km"].fillna(0)
    df["assigned_officer"] = df["assigned_officer"].fillna("Unknown")

    # Select features for prediction
    feature_columns = [
        "category",
        "type",
        "state",
        "city",
        "assigned_officer",
        "asset_age_years",
        "maintenance_frequency",
        "days_to_maintenance",
        "usage_hours",
        "usage_intensity",
        "mileage_km",
        "annual_mileage",
        "age_category",
        "high_risk_location",
        "critical_asset",
        "maintenance_overdue",
        "high_usage_hours",
        "high_mileage",
        "has_ammunition_data",
        "ammunition_count",
        "warranty_expired",
    ]

    return df[feature_columns]


def generate_ai_report(report_config, kpis, df):
    """Generate AI-powered report using OpenAI."""
    try:
        # Initialize OpenAI client
        client = openai.OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

        # Prepare data summary for the AI
        data_summary = f"""
        PDRM Asset & Facility Monitoring Report Data:
        
        KEY PERFORMANCE INDICATORS:
        - Total Assets: {kpis['total_assets']:,}
        - Assets in Good Condition: {kpis['good_assets']:,} ({kpis['good_pct']:.1f}%)
        - Assets Needing Action: {kpis['needs_action']:,} ({kpis['needs_action_pct']:.1f}%)
        - Damaged Assets: {kpis['damaged_assets']:,} ({kpis['damaged_pct']:.1f}%)
        - Assets with Overdue Maintenance: {kpis['overdue_maintenance']:,}
        
        ASSET DISTRIBUTION:
        {df.groupby(['category', 'condition']).size().to_string()}
        
        STATE-WISE DISTRIBUTION:
        {df.groupby(['state', 'condition']).size().head(20).to_string()}
        """

        prompt = f"""
        You are a professional asset management analyst for the Royal Malaysia Police (PDRM). 
        Generate a comprehensive executive report based on the following data:
        
        {data_summary}
        
        Please include the following sections in your report:
        1. EXECUTIVE SUMMARY
        2. KEY FINDINGS AND INSIGHTS
        3. ASSET CONDITION ANALYSIS
        4. MAINTENANCE RECOMMENDATIONS
        5. RISK ASSESSMENT
        6. STRATEGIC RECOMMENDATIONS
        
        The report should be professional, actionable, and suitable for senior management decision-making.
        Focus on operational efficiency, cost optimization, and asset reliability.
        """

        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {
                    "role": "system",
                    "content": "You are a professional asset management analyst.",
                },
                {"role": "user", "content": prompt},
            ],
            max_tokens=2000,
            temperature=0.7,
        )

        return response.choices[0].message.content

    except Exception as e:
        return f"Error generating AI report: {str(e)}"


def generate_expired_stock_report(df, filename="Expired_Stock_List.docx"):
    """Generate Expired Stock List report based on Malaysian government template KEW.PS-6."""
    try:
        doc = Document()

        # Header section
        header = doc.add_paragraph()
        header.alignment = 1  # Center alignment
        header_run = header.add_run("Polis Diraja Malaysia")
        header_run.bold = True
        header_run.font.size = Pt(12)

        doc.add_paragraph()  # Empty line

        # Form number - top right
        form_para = doc.add_paragraph()
        form_para.alignment = 2  # Right alignment
        form_run = form_para.add_run("KEW.PS-6")
        form_run.bold = True
        form_run.font.size = Pt(10)

        doc.add_paragraph()  # Empty line

        # Title
        title = doc.add_heading("SENARAI STOK BERTARIKH LUPUT", level=1)
        title.alignment = 1  # Center alignment

        doc.add_paragraph()  # Empty line

        # Information fields (left aligned)
        info_para1 = doc.add_paragraph()
        info_para1.add_run("Kementerian/Jabatan: ").bold = True
        info_para1.add_run("Jabatan Logistik & Asset")

        info_para2 = doc.add_paragraph()
        info_para2.add_run("Kategori Stor: ").bold = True
        info_para2.add_run("Aset & Kemudahan Kerajaan")

        info_para3 = doc.add_paragraph()
        info_para3.add_run("Bulan: ").bold = True
        months_bm = {
            "January": "Januari",
            "February": "Februari",
            "March": "Mac",
            "April": "April",
            "May": "Mei",
            "June": "Jun",
            "July": "Julai",
            "August": "Ogos",
            "September": "September",
            "October": "Oktober",
            "November": "November",
            "December": "Disember",
        }
        current_month_en = datetime.now().strftime("%B")
        current_month_bm = months_bm.get(current_month_en, current_month_en)
        info_para3.add_run(f"{current_month_bm} {datetime.now().year}")

        doc.add_paragraph()  # Empty line

        # Create table - match exact format from image
        # Calculate items that need attention (damaged or overdue maintenance)
        expired_items = df[
            (df["condition"] == "Damaged")
            | (
                pd.to_datetime(df["next_maintenance_due"], errors="coerce")
                < datetime.now()
            )
        ].copy()

        # Create table with exact column structure from reference image
        # Structure: Bil, Perihal Stok, No.Kod, Lokasi Stok, Kuantiti Termasuk, Tarikh Luput, 6 bulan, 5 bulan, 4 bulan, 3 bulan, 2 bulan, 1 bulan, Catatan
        table = doc.add_table(rows=2, cols=13)  # Header + sub-header rows
        table.style = "Table Grid"

        # Set column widths to match reference format
        col_widths = [
            Inches(0.3),  # Bil.
            Inches(1.5),  # Perihal Stok
            Inches(0.7),  # No. Kod
            Inches(1.0),  # Lokasi Stok
            Inches(0.6),  # Kuantiti Termasuk
            Inches(0.7),  # Tarikh Luput
            Inches(0.4),  # 6 bulan
            Inches(0.4),  # 5 bulan
            Inches(0.4),  # 4 bulan
            Inches(0.4),  # 3 bulan
            Inches(0.4),  # 2 bulan
            Inches(0.4),  # 1 bulan
            Inches(1.0),  # Catatan
        ]

        for i, col in enumerate(table.columns):
            if i < len(col_widths):
                for cell in col.cells:
                    cell.width = col_widths[i]

        # Main header row
        hdr_cells = table.rows[0].cells

        # Set main headers according to reference image
        hdr_cells[0].text = "Bil."
        hdr_cells[1].text = "Perihal Stok"
        hdr_cells[2].text = "No. Kod"
        hdr_cells[3].text = "Lokasi Stok"
        hdr_cells[4].text = "Kuantiti Termasuk"
        hdr_cells[5].text = "Tarikh Luput"

        # Merge cells for "Baki Stok" header (columns 6-11) - this creates one single merged cell
        merged_cell = (
            hdr_cells[6]
            .merge(hdr_cells[7])
            .merge(hdr_cells[8])
            .merge(hdr_cells[9])
            .merge(hdr_cells[10])
            .merge(hdr_cells[11])
        )
        merged_cell.text = "Baki Stok Bagi 6 Bulan Sebelum Tarikh Tempoh Luput\n*(Kuantiti dan tarikh kemaskini)"

        hdr_cells[12].text = "Catatan"

        # Format main header
        for i, cell in enumerate(hdr_cells):
            # Skip the merged cells (7-11) as they are now part of cell 6
            if i in [7, 8, 9, 10, 11]:
                continue

            for paragraph in cell.paragraphs:
                paragraph.alignment = 1  # Center alignment
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = (
                        Pt(9) if i != 6 else Pt(8)
                    )  # Smaller font for merged header

        # Sub-header row for month columns
        sub_hdr_cells = table.rows[1].cells
        sub_hdr_cells[0].text = ""  # Bil.
        sub_hdr_cells[1].text = ""  # Perihal Stok
        sub_hdr_cells[2].text = ""  # No. Kod
        sub_hdr_cells[3].text = ""  # Lokasi
        sub_hdr_cells[4].text = ""  # Kuantiti
        sub_hdr_cells[5].text = ""  # Tarikh Luput
        sub_hdr_cells[6].text = "6\nbulan"
        sub_hdr_cells[7].text = "5\nbulan"
        sub_hdr_cells[8].text = "4\nbulan"
        sub_hdr_cells[9].text = "3\nbulan"
        sub_hdr_cells[10].text = "2\nbulan"
        sub_hdr_cells[11].text = "1\nbulan"
        sub_hdr_cells[12].text = ""  # Catatan

        # Format sub-headers
        for i, cell in enumerate(sub_hdr_cells):
            for paragraph in cell.paragraphs:
                paragraph.alignment = 1  # Center alignment
                for run in paragraph.runs:
                    run.font.size = Pt(8)
                    if 6 <= i <= 11:  # Month columns
                        run.font.bold = True

        # Add data rows (limit to first 10 items for readability)
        for idx, (_, item) in enumerate(expired_items.head(10).iterrows(), 1):
            row_cells = table.add_row().cells

            # Determine status and dates
            if item["condition"] == "Damaged":
                tarikh_luput = "Rosak"
                catatan = f"Aset rosak - {item['condition']}"
            else:
                try:
                    next_maintenance = pd.to_datetime(item["next_maintenance_due"])
                    if pd.notna(next_maintenance):
                        tarikh_luput = next_maintenance.strftime("%d/%m/%Y")
                        catatan = "Penyelenggaraan tertunggak"
                    else:
                        tarikh_luput = "Tiada data"
                        catatan = "Tiada jadual penyelenggaraan"
                except:
                    tarikh_luput = "Tiada data"
                    catatan = "Tarikh tidak sah"

            # Asset age for quantity distribution
            try:
                asset_age = (
                    datetime.now() - pd.to_datetime(item["purchase_date"])
                ).days / 365.25
            except:
                asset_age = 0

            # Sample quantity distribution for each month column
            qty_6_month = "1" if asset_age > 5 else ""
            qty_5_month = "1" if 4 < asset_age <= 5 else ""
            qty_4_month = "1" if 3 < asset_age <= 4 else ""
            qty_3_month = "1" if 2 < asset_age <= 3 else ""
            qty_2_month = "1" if 1 < asset_age <= 2 else ""
            qty_1_month = "1" if asset_age <= 1 else ""

            # Fill data - now with 13 columns structure
            data = [
                str(idx),  # Bil
                f"{item['type']}\n{item['vendor'] if pd.notna(item['vendor']) else 'N/A'}".strip(),  # Perihal Stok
                item["asset_id"],  # No. Kod
                (  # Lokasi Stok
                    f"{item['city']}, {item['state']}"
                    if pd.notna(item["city"]) and pd.notna(item["state"])
                    else "N/A"
                ),
                "1",  # Kuantiti Termasuk - Assuming 1 unit per asset
                tarikh_luput,  # Tarikh Luput
                qty_6_month,  # 6 bulan
                qty_5_month,  # 5 bulan
                qty_4_month,  # 4 bulan
                qty_3_month,  # 3 bulan
                qty_2_month,  # 2 bulan
                qty_1_month,  # 1 bulan
                catatan,  # Catatan
            ]

            # Fill the row cells
            for i, cell_data in enumerate(data):
                if i < 13:  # Fill all 13 columns
                    row_cells[i].text = str(cell_data)
                    # Format cell text
                    for paragraph in row_cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(8)

        # Add empty rows for manual completion
        for i in range(15):  # Add 15 empty rows
            row_cells = table.add_row().cells
            for j in range(13):  # Now 13 columns
                row_cells[j].text = ""

        doc.add_paragraph()  # Empty line

        # Footer note in Bahasa Malaysia
        footer_note = doc.add_paragraph()
        footer_note.add_run(
            "* Dikemaskini kuantiti dan tarikh luput berdasarkan keadaan aset dan jadual penyelenggaraan dalam bulan yang sama."
        )
        footer_run = footer_note.runs[0]
        footer_run.font.size = Pt(8)

        # Save to BytesIO
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)

        return doc_io.getvalue()

    except Exception as e:
        st.error(f"Ralat menghasilkan laporan stok luput: {str(e)}")
        return None


def render_ai_chatbot(df):
    """Render AI Assistant chatbot with pre-defined prompts."""
    st.markdown("---")
    st.markdown(
        '<div class="subsection-header">🤖 AI Assistant - Dashboard Intelligence</div>',
        unsafe_allow_html=True,
    )

    st.markdown(
        """
    Ask the AI Assistant about PDRM asset data to get instant analysis and insights.
    """
    )

    # Pre-defined prompt questions
    col1, col2, col3 = st.columns(3)

    with col1:
        if st.button("📊 Overall Asset Status", use_container_width=True):
            st.session_state["ai_query"] = (
                "Provide a comprehensive summary of overall PDRM asset status based on the dashboard data"
            )

        if st.button("🔧 Assets Need Maintenance", use_container_width=True):
            st.session_state["ai_query"] = (
                "List assets that require immediate maintenance and recommended actions"
            )

    with col2:
        if st.button("📈 Asset Trend Analysis", use_container_width=True):
            st.session_state["ai_query"] = (
                "Analyze purchasing trends and asset conditions by year and state"
            )

        if st.button("⚠️ High Risk Assets", use_container_width=True):
            st.session_state["ai_query"] = (
                "Identify high-risk assets and replacement priorities"
            )

    with col3:
        if st.button("💰 Operational Cost Analysis", use_container_width=True):
            st.session_state["ai_query"] = (
                "Provide operational cost analysis and asset optimization recommendations"
            )

        if st.button("📋 KPI Performance Report", use_container_width=True):
            st.session_state["ai_query"] = (
                "Analyze asset KPI performance and improvement recommendations"
            )

    # Chat interface
    st.markdown("### 💬 Chat with AI Assistant")

    # Initialize session state for chat history
    if "chat_history" not in st.session_state:
        st.session_state.chat_history = []

    # Text input for custom questions
    user_question = st.text_input(
        "Ask your question about PDRM asset data:",
        value=st.session_state.get("ai_query", ""),
        placeholder="Example: How many damaged assets are there in Selangor?",
    )

    col_send, col_clear = st.columns([1, 4])
    with col_send:
        send_pressed = st.button("📤 Send", type="primary")
    with col_clear:
        if st.button("🗑️ Clear Chat"):
            st.session_state.chat_history = []
            st.session_state["ai_query"] = ""
            st.rerun()

    # Process user question
    if send_pressed and user_question:
        # Clear the pre-selected query
        if "ai_query" in st.session_state:
            del st.session_state["ai_query"]

        # Add user message to chat history
        st.session_state.chat_history.append({"role": "user", "content": user_question})

        with st.spinner("AI is analyzing the data..."):
            # Generate AI response based on the dashboard data
            ai_response = generate_ai_response(user_question, df)

            # Add AI response to chat history
            st.session_state.chat_history.append(
                {"role": "assistant", "content": ai_response}
            )

        st.rerun()

    # Display chat history
    if st.session_state.chat_history:
        st.markdown("### 💬 Chat History")

        # Create a container for chat messages
        chat_container = st.container()

        with chat_container:
            for i, message in enumerate(st.session_state.chat_history):
                if message["role"] == "user":
                    st.markdown(
                        f"""
                    <div class="chat-message-user">
                        <strong>👤 You:</strong><br>
                        <span>{message["content"]}</span>
                    </div>
                    """,
                        unsafe_allow_html=True,
                    )
                else:
                    st.markdown(
                        f"""
                    <div class="chat-message-assistant">
                        <strong>🤖 AI Assistant:</strong><br>
                        <span>{message["content"]}</span>
                    </div>
                    """,
                        unsafe_allow_html=True,
                    )


def generate_ai_response(question, df):
    """Generate AI response using OpenAI API based on dashboard data and user question."""
    try:
        # Initialize OpenAI client
        client = openai.OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

        # Analyze the dataset to provide relevant statistics
        total_assets = len(df)
        damaged_assets = len(df[df["condition"] == "Damaged"])
        good_condition = len(df[df["condition"] == "Good"])
        needs_action = len(df[df["condition"] == "Needs Action"])

        # State distribution
        state_counts = df["state"].value_counts()
        top_states = state_counts.head(5)

        # Category distribution
        category_counts = df["category"].value_counts()
        top_categories = category_counts.head(5)

        # Overdue maintenance
        current_date = datetime.now()
        overdue_maintenance = len(
            df[
                pd.to_datetime(df["next_maintenance_due"], errors="coerce")
                < current_date
            ]
        )

        # Purchase year trends
        df_copy = df.copy()
        df_copy["purchase_year"] = pd.to_datetime(df_copy["purchase_date"]).dt.year
        year_counts = df_copy["purchase_year"].value_counts().sort_index()
        recent_years = year_counts.tail(5)

        # Create comprehensive data context
        data_context = f"""
PDRM (Royal Malaysia Police) Assets Database Analysis:

OVERALL STATISTICS:
- Total Assets: {total_assets:,} units
- Assets in Good Condition: {good_condition:,} ({good_condition/total_assets*100:.1f}%)
- Damaged Assets: {damaged_assets:,} ({damaged_assets/total_assets*100:.1f}%)  
- Assets Needing Action: {needs_action:,} ({needs_action/total_assets*100:.1f}%)
- Assets with Overdue Maintenance: {overdue_maintenance:,} ({overdue_maintenance/total_assets*100:.1f}%)

TOP 5 STATES BY ASSET COUNT:
{chr(10).join([f"- {state}: {count:,} assets ({count/total_assets*100:.1f}%)" for state, count in top_states.items()])}

TOP 5 ASSET CATEGORIES:
{chr(10).join([f"- {category}: {count:,} units ({count/total_assets*100:.1f}%)" for category, count in top_categories.items()])}

RECENT PURCHASE TRENDS (Last 5 Years):
{chr(10).join([f"- Year {year}: {count:,} assets purchased" for year, count in recent_years.items()])}

MAINTENANCE INSIGHTS:
- Maintenance Compliance Rate: {(total_assets-overdue_maintenance)/total_assets*100:.1f}%
- Assets Requiring Immediate Attention: {damaged_assets + needs_action:,} units
- Operational Efficiency: {good_condition/total_assets*100:.1f}% of assets fully operational
"""

        # Create system prompt for PDRM context
        system_prompt = """You are an AI assistant specialized in PDRM (Royal Malaysia Police) asset management and analysis. 
You have access to comprehensive asset data and should provide accurate, actionable insights based on the data provided.

Your responses should be:
- Professional and suitable for law enforcement management
- Based strictly on the data provided
- Include specific numbers and percentages from the dataset
- Provide actionable recommendations
- Use appropriate emojis for better readability
- Format responses in clear markdown with sections

Focus on operational efficiency, maintenance priorities, budget optimization, and strategic asset management for law enforcement operations."""

        # Create user prompt with question and data
        user_prompt = f"""Based on the PDRM asset data provided below, please answer this question: "{question}"

{data_context}

Please provide a comprehensive analysis based strictly on this data, including relevant statistics, insights, and actionable recommendations for PDRM asset management."""

        # Call OpenAI API
        response = client.chat.completions.create(
            model="gpt-4o-mini",  # Using GPT-4o-mini for cost efficiency
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            max_tokens=1500,
            temperature=0.3,  # Lower temperature for more factual responses
            top_p=1,
            frequency_penalty=0,
            presence_penalty=0,
        )

        ai_response = response.choices[0].message.content

        # Add footer with data source information
        ai_response += f"""

---
� **Data Source:** PDRM Asset Management System | **Last Updated:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} | **Total Records:** {total_assets:,}
"""

        return ai_response

    except Exception as e:
        # Fallback response if OpenAI API fails
        error_msg = str(e)

        # Check if it's an API key issue
        if "api_key" in error_msg.lower() or "authentication" in error_msg.lower():
            return """
❌ **OpenAI API Configuration Error**

The AI assistant requires a valid OpenAI API key to function properly.

**To fix this:**
1. Obtain an OpenAI API key from https://platform.openai.com/
2. Set it as an environment variable: `OPENAI_API_KEY=your_key_here`
3. Restart the dashboard

**For now, here's a basic data summary:**

📊 **Current PDRM Asset Status:**
- Total assets in database: {len(df):,} units
- Assets requiring attention: {len(df[df["condition"].isin(["Damaged", "Needs Action"])]):,} units

**Contact IT Support:** asset.management@pdrm.gov.my
"""
        else:
            return f"""
❌ **AI Assistant Temporarily Unavailable**

There was an error connecting to the AI service: {error_msg}

**Basic Asset Summary:**
- Total Assets: {len(df):,} units
- Condition Breakdown:
  - Good: {len(df[df["condition"] == "Good"]):,} units
  - Damaged: {len(df[df["condition"] == "Damaged"]):,} units  
  - Needs Action: {len(df[df["condition"] == "Needs Action"]):,} units

**Please try again or contact technical support.**
**Support:** asset.management@pdrm.gov.my
"""


def export_report_to_word_with_charts(
    report_content, charts_data, df, filename="PDRM_Asset_Report.docx"
):
    """Export comprehensive report to Word document with charts."""
    try:
        doc = Document()

        # Add title
        title = doc.add_heading("PDRM Asset & Facility Monitoring Report", 0)
        title.alignment = 1  # Center alignment

        # Add metadata
        doc.add_paragraph(
            f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        )
        doc.add_paragraph("Royal Malaysia Police - Asset Management Division")
        doc.add_page_break()

        # Add executive summary
        doc.add_heading("EXECUTIVE SUMMARY", level=1)

        # Add KPIs table
        doc.add_heading("Key Performance Indicators", level=2)
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"

        # Add table headers
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Metric"
        hdr_cells[1].text = "Value"

        # Add KPI data
        kpi_data = [
            ("Total Assets", f"{charts_data['total_assets']:,}"),
            (
                "Assets in Good Condition",
                f"{charts_data['good_assets']:,} ({charts_data['good_pct']:.1f}%)",
            ),
            (
                "Assets Needing Action",
                f"{charts_data['needs_action']:,} ({charts_data['needs_action_pct']:.1f}%)",
            ),
            (
                "Damaged Assets",
                f"{charts_data['damaged_assets']:,} ({charts_data['damaged_pct']:.1f}%)",
            ),
            ("Overdue Maintenance", f"{charts_data['overdue_maintenance']:,}"),
        ]

        for metric, value in kpi_data:
            row_cells = table.add_row().cells
            row_cells[0].text = metric
            row_cells[1].text = value

        doc.add_paragraph()

        # Generate and add charts
        doc.add_heading("VISUAL ANALYTICS", level=1)

        # 1. Asset Distribution Chart
        doc.add_heading("Asset Distribution by Category and Condition", level=2)

        try:
            # Create distribution chart
            type_condition = (
                df.groupby(["category", "condition"]).size().unstack(fill_value=0)
            )
            fig_dist = px.bar(
                type_condition.reset_index().melt(
                    id_vars=["category"], var_name="Condition", value_name="Count"
                ),
                x="category",
                y="Count",
                color="Condition",
                title="Asset Distribution by Category and Condition",
                labels={"category": "Asset Category", "Count": "Number of Assets"},
                color_discrete_map={
                    "Good": "#2E8B57",
                    "Needs Action": "#FFD700",
                    "Damaged": "#DC143C",
                },
            )
            fig_dist.update_layout(height=400, title_font_size=14, title_x=0.5)

            # Save chart as image safely
            try:
                img_data = safe_chart_to_image(fig_dist, width=800, height=400)
                if img_data:
                    img_stream = io.BytesIO(img_data)
                    doc.add_picture(img_stream, width=Inches(6))
                else:
                    doc.add_paragraph(
                        "Chart image not available (Kaleido dependency issue)"
                    )
            except Exception as e:
                doc.add_paragraph(f"Chart image could not be generated: {str(e)}")
            doc.add_paragraph()

        except Exception as e:
            doc.add_paragraph(f"Chart generation error: {str(e)}")

        # 2. Asset Condition Pie Chart
        doc.add_heading("Overall Asset Condition Distribution", level=2)

        try:
            condition_counts = df["condition"].value_counts()
            fig_pie = px.pie(
                values=condition_counts.values,
                names=condition_counts.index,
                title="Overall Asset Condition Distribution",
                color_discrete_map={
                    "Good": "#2E8B57",
                    "Needs Action": "#FFD700",
                    "Damaged": "#DC143C",
                },
            )
            fig_pie.update_layout(height=400, title_font_size=14, title_x=0.5)

            # Save chart as image safely
            try:
                img_data = safe_chart_to_image(fig_pie, width=800, height=400)
                if img_data:
                    img_stream = io.BytesIO(img_data)
                    doc.add_picture(img_stream, width=Inches(6))
                else:
                    doc.add_paragraph(
                        "Chart image not available (Kaleido dependency issue)"
                    )
            except Exception as e:
                doc.add_paragraph(f"Chart image could not be generated: {str(e)}")
            doc.add_paragraph()

        except Exception as e:
            doc.add_paragraph(f"Chart generation error: {str(e)}")

        # 3. State-wise Heatmap (simplified table for Word)
        doc.add_heading("Asset Distribution by State", level=2)

        try:
            state_condition = (
                df.groupby(["state", "condition"]).size().unstack(fill_value=0)
            )

            # Create table for state data
            state_table = doc.add_table(rows=1, cols=4)
            state_table.style = "Table Grid"

            # Headers
            headers = state_table.rows[0].cells
            headers[0].text = "State"
            headers[1].text = "Good"
            headers[2].text = "Needs Action"
            headers[3].text = "Damaged"

            # Add top 10 states data
            for state in state_condition.head(10).index:
                row_cells = state_table.add_row().cells
                row_cells[0].text = state
                row_cells[1].text = str(state_condition.loc[state].get("Good", 0))
                row_cells[2].text = str(
                    state_condition.loc[state].get("Needs Action", 0)
                )
                row_cells[3].text = str(state_condition.loc[state].get("Damaged", 0))

            doc.add_paragraph()

        except Exception as e:
            doc.add_paragraph(f"Table generation error: {str(e)}")

        # Add AI-generated content
        doc.add_heading("DETAILED ANALYSIS", level=1)

        # Split AI content into paragraphs
        paragraphs = report_content.split("\n\n")
        for paragraph in paragraphs:
            if paragraph.strip():
                # Check if it's a heading (starts with number or all caps)
                if paragraph.strip().isupper() or (
                    paragraph.strip() and paragraph.strip()[0].isdigit()
                ):
                    doc.add_heading(paragraph.strip(), level=2)
                else:
                    doc.add_paragraph(paragraph.strip())

        # Add category breakdown section
        doc.add_heading("ASSET CATEGORY BREAKDOWN", level=1)

        categories = ["Weapons", "ICT Assets", "Vehicles", "Devices"]
        for category in categories:
            if category in charts_data["category_stats"]:
                stats = charts_data["category_stats"][category]
                doc.add_heading(f"{category}", level=2)

                # Add category statistics
                doc.add_paragraph(f"Total Assets: {stats['total_assets']:,}")
                doc.add_paragraph(
                    f"Good Condition: {stats['good_count']:,} ({stats['good_pct']:.1f}%)"
                )
                doc.add_paragraph(
                    f"Needs Action: {stats['needs_action_count']:,} ({stats['needs_action_pct']:.1f}%)"
                )
                doc.add_paragraph(
                    f"Damaged: {stats['damaged_count']:,} ({stats['damaged_pct']:.1f}%)"
                )
                doc.add_paragraph(f"Average Age: {stats['avg_age_years']:.1f} years")

                # Add top asset types
                doc.add_paragraph("Top Asset Types:")
                for asset_type, count in stats["top_types"].head(5).items():
                    doc.add_paragraph(
                        f"• {asset_type}: {count:,} units", style="List Bullet"
                    )

                # Add category-specific chart
                try:
                    category_df = df[df["category"] == category]
                    type_condition_cat = (
                        category_df.groupby(["type", "condition"])
                        .size()
                        .reset_index(name="count")
                    )

                    if not type_condition_cat.empty:
                        fig_cat = px.bar(
                            type_condition_cat,
                            x="type",
                            y="count",
                            color="condition",
                            title=f"{category} - Asset Types and Conditions",
                            labels={
                                "type": "Asset Type",
                                "count": "Number of Assets",
                                "condition": "Condition",
                            },
                            color_discrete_map={
                                "Good": "#2E8B57",
                                "Needs Action": "#FFD700",
                                "Damaged": "#DC143C",
                            },
                        )
                        fig_cat.update_layout(
                            height=300,
                            title_font_size=12,
                            title_x=0.5,
                            xaxis={"tickangle": 45},
                        )

                        # Save category chart as image
                        img_data = safe_chart_to_image(fig_cat, width=800, height=300)
                        if img_data:
                            img_stream = io.BytesIO(img_data)
                            doc.add_picture(img_stream, width=Inches(6))
                        else:
                            doc.add_paragraph(
                                "Chart image not available (Kaleido dependency issue)"
                            )

                except Exception as e:
                    doc.add_paragraph(f"Category chart error: {str(e)}")

                doc.add_paragraph()

        # Add recommendations
        doc.add_heading("STRATEGIC RECOMMENDATIONS", level=1)

        recommendations = [
            "Prioritize maintenance for overdue assets to prevent further deterioration",
            "Implement predictive maintenance schedule based on asset age and usage patterns",
            "Consider replacement program for assets with consistently poor condition ratings",
            "Enhance monitoring systems for high-risk asset categories",
            "Establish preventive maintenance protocols for critical equipment",
            "Allocate budget for high-maintenance categories (Weapons and Vehicles)",
            "Implement asset tracking system for better maintenance scheduling",
            "Train staff on proper asset care and maintenance procedures",
        ]

        for rec in recommendations:
            doc.add_paragraph(rec, style="List Bullet")

        # Add footer
        doc.add_page_break()
        doc.add_heading("REPORT METADATA", level=1)
        doc.add_paragraph(
            f"Report Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        )
        doc.add_paragraph(f"Total Records Analyzed: {len(df):,}")
        doc.add_paragraph(
            f"Data Period: {df['purchase_date'].min()} to {df['purchase_date'].max()}"
        )
        doc.add_paragraph("Generated by: PDRM Asset Management Dashboard")
        doc.add_paragraph("Classification: Internal Use")

        # Save to memory
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)

        return doc_buffer.getvalue()

    except Exception as e:
        st.error(f"Error creating Word document: {str(e)}")
        return None


def export_report_to_pdf(report_content, filename="PDRM_Asset_Report.pdf"):
    """Export report to PDF."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)

    # Add title
    pdf.set_font("Arial", "B", 16)
    pdf.cell(0, 10, "PDRM Asset & Facility Monitoring Report", ln=True, align="C")
    pdf.ln(10)

    # Add report content
    pdf.set_font("Arial", size=10)
    lines = report_content.split("\n")
    for line in lines:
        if len(line) > 80:
            # Split long lines
            words = line.split(" ")
            current_line = ""
            for word in words:
                if len(current_line + word) < 80:
                    current_line += word + " "
                else:
                    pdf.cell(
                        0,
                        6,
                        current_line.encode("latin-1", "replace").decode("latin-1"),
                        ln=True,
                    )
                    current_line = word + " "
            if current_line:
                pdf.cell(
                    0,
                    6,
                    current_line.encode("latin-1", "replace").decode("latin-1"),
                    ln=True,
                )
        else:
            pdf.cell(0, 6, line.encode("latin-1", "replace").decode("latin-1"), ln=True)

    return pdf.output(dest="S").encode("latin-1")


def get_logo_base64():
    """Convert PDRM logo to base64 for embedding in HTML."""
    try:
        import base64

        with open("Royal_Malaysian_Police.jpg", "rb") as image_file:
            encoded_string = base64.b64encode(image_file.read()).decode()
            return encoded_string
    except:
        # Return empty string if logo file not found
        return ""


def main():
    """Main dashboard application."""

    # Main header with logo inside card box
    logo_base64 = get_logo_base64()

    if logo_base64:
        # Display header with logo
        colors = get_theme_colors()
        st.markdown(
            f"""
            <div style="background: {colors['bg_gradient']}; 
                        padding: 2rem; 
                        border-radius: 15px; 
                        box-shadow: {colors['shadow']};
                        border: 1px solid {colors['border']};
                        margin-bottom: 2rem;
                        display: flex;
                        align-items: center;
                        justify-content: center;
                        gap: 2rem;">
                <div style="flex-shrink: 0;">
                    <img src="data:image/jpeg;base64,{logo_base64}" 
                         style="width: 80px; height: 80px; border-radius: 8px; object-fit: contain;" 
                         alt="PDRM Logo">
                </div>
                <div style="text-align: center;">
                    <div style="font-size: 2.5rem; font-weight: bold; color: white;">
                        Jabatan Logistik & Teknologi
                    </div>
                </div>
            </div>
            """,
            unsafe_allow_html=True,
        )
    else:
        # Fallback without logo
        st.markdown(
            f"""
            <div style="background: {colors['bg_gradient']}; 
                        padding: 2rem; 
                        border-radius: 15px; 
                        box-shadow: {colors['shadow']};
                        border: 1px solid {colors['border']};
                        margin-bottom: 2rem;
                        text-align: center;">
                <div style="font-size: 2.5rem; font-weight: bold; color: {colors['text_primary']};">
                    🏛️ Jabatan Logistik & Asset
                </div>
            </div>
            """,
            unsafe_allow_html=True,
        )

    # Load data
    df = load_data()
    if df is None:
        st.error("Unable to load data. Please check if 'assets_data.csv' exists.")
        return

    # Sidebar filters with dark theme styling
    st.sidebar.markdown(
        """
        <div style="
            background: linear-gradient(135deg, #1e293b 0%, #334155 100%);
            padding: 1rem;
            border-radius: 10px;
            border: 1px solid #475569;
            margin-bottom: 1.5rem;
            text-align: center;
            box-shadow: 0 4px 12px rgba(0,0,0,0.3);
        ">
            <div style="color: #60a5fa; font-size: 1.1rem; font-weight: bold;">
                🔍 Filters
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    # Asset type filter
    asset_types = ["All"] + list(df["category"].unique())
    selected_type = st.sidebar.selectbox("Asset Type", asset_types, index=0)

    # State filter
    states = ["All"] + list(df["state"].unique())
    selected_state = st.sidebar.selectbox("State", states, index=0)

    # Year range filter with dropdown
    st.sidebar.markdown(
        """
        <div style="color: white; font-weight: bold; margin-bottom: 0.5rem; margin-top: 1rem;">
            📅 Purchase Year Range
        </div>
        """,
        unsafe_allow_html=True,
    )
    min_year = int(df["purchase_date"].dt.year.min())
    max_year = int(df["purchase_date"].dt.year.max())

    # Create dropdown for year range selection
    col1, col2 = st.sidebar.columns(2)
    with col1:
        start_year = st.selectbox("From Year", range(min_year, max_year + 1), index=0)
    with col2:
        end_year = st.selectbox(
            "To Year", range(min_year, max_year + 1), index=max_year - min_year
        )

    selected_years = (start_year, end_year)

    # Apply filters
    filtered_df = df.copy()
    if selected_type != "All":
        filtered_df = filtered_df[filtered_df["category"] == selected_type]
    if selected_state != "All":
        filtered_df = filtered_df[filtered_df["state"] == selected_state]

    filtered_df = filtered_df[
        (filtered_df["purchase_date"].dt.year >= selected_years[0])
        & (filtered_df["purchase_date"].dt.year <= selected_years[1])
    ]

    # Calculate KPIs
    kpis = calculate_kpis(filtered_df)

    # Create tabs
    tab1, tab2, tab3 = st.tabs(
        [
            "📊 Descriptive Analytics",
            "🔮 Predictive Analytics",
            "📝 AI Report Generator",
        ]
    )

    with tab1:
        st.markdown(
            '<div class="section-header">📊 Descriptive Analytics</div>',
            unsafe_allow_html=True,
        )

        # Display KPIs
        display_kpis(kpis, filtered_df)

        # Charts
        col1, col2 = st.columns(2)

        with col1:
            # Trend analysis
            fig_trend = create_trend_analysis(filtered_df)
            fig_trend.update_layout(height=500)  # Set consistent height
            st.plotly_chart(fig_trend, use_container_width=True)

            # Simple maintenance timeline
            today = datetime.now()
            upcoming_maintenance = filtered_df[
                (filtered_df["next_maintenance_due"] >= today)
                & (filtered_df["next_maintenance_due"] <= today + timedelta(days=90))
            ]

            fig_upcoming = px.histogram(
                upcoming_maintenance,
                x="category",
                title="Assets Due for Maintenance (Next 3 Months)",
                labels={"category": "Category", "count": "Number of Assets"},
                color_discrete_sequence=[
                    "#475569"
                ],  # Changed from #87CEEB to match dark theme
            )
            fig_upcoming = apply_chart_styling(fig_upcoming, height=500, title_size=14)
            st.plotly_chart(fig_upcoming, use_container_width=True)

        with col2:
            # Distribution chart
            fig_dist = create_distribution_chart(filtered_df)
            fig_dist.update_layout(height=500)  # Set consistent height
            st.plotly_chart(fig_dist, use_container_width=True)

            # Asset condition distribution by state stacked bar chart
            fig_state_condition = create_state_condition_chart(filtered_df)
            fig_state_condition.update_layout(height=500)  # Set consistent height
            st.plotly_chart(fig_state_condition, use_container_width=True)

        # Detailed Category Breakdown Section
        st.markdown(
            '<div class="subsection-header">🔍 Detailed Category Analysis</div>',
            unsafe_allow_html=True,
        )

        # Category summary metrics
        category_metrics = create_category_summary_metrics(filtered_df)

        # Display category metrics in columns with theme-aware cardboxes
        cols = st.columns(len(category_metrics))
        colors = get_theme_colors()
        for i, (category, metrics) in enumerate(category_metrics.items()):
            with cols[i]:
                st.markdown(
                    f"""
                    <div style="
                        background: {colors['bg_gradient']};
                        padding: 1.5rem;
                        border-radius: 12px;
                        border: 1px solid {colors['border']};
                        box-shadow: 0 4px 15px rgba(0,0,0,{'0.4' if colors['text_primary'] == 'white' else '0.1'});
                        margin-bottom: 1rem;
                        text-align: center;
                        color: {colors['text_primary']};
                    ">
                        <div style="font-size: 2rem; font-weight: bold; margin-bottom: 0.5rem;">
                            {metrics['total_assets']:,}
                        </div>
                        <div style="font-size: 1.1rem; font-weight: 600; margin-bottom: 1rem; opacity: 0.9;">
                            {category} Total
                        </div>
                        <div style="display: flex; flex-direction: column; gap: 0.3rem; font-size: 0.9rem;">
                            <div style="color: #22c55e;">✅ Good: {metrics['good_pct']:.1f}%</div>
                            <div style="color: #f59e0b;">⚠️ Needs Action: {metrics['needs_action_pct']:.1f}%</div>
                            <div style="color: #ef4444;">❌ Damaged: {metrics['damaged_pct']:.1f}%</div>
                        </div>
                    </div>
                """,
                    unsafe_allow_html=True,
                )

        # Detailed breakdown charts for each category
        st.markdown(
            '<div class="subsection-header">📊 Asset Type Breakdown by Category</div>',
            unsafe_allow_html=True,
        )
        category_charts = create_detailed_category_breakdown(filtered_df)

        # Create tabs for each category
        if len(category_charts) > 0:
            category_tabs = st.tabs(list(category_charts.keys()))
            for i, (category, fig) in enumerate(category_charts.items()):
                with category_tabs[i]:
                    # Set consistent height for category charts
                    fig.update_layout(height=500)
                    st.plotly_chart(fig, use_container_width=True)

                    # Show top asset types for this category
                    category_df = filtered_df[filtered_df["category"] == category]
                    top_types = category_df["type"].value_counts().head(5)

                    st.write("**Top Asset Types:**")
                    for asset_type, count in top_types.items():
                        percentage = (count / len(category_df)) * 100
                        st.write(
                            f"• {asset_type}: {count:,} assets ({percentage:.1f}%)"
                        )

        # Maintenance Priority Analysis
        st.markdown(
            '<div class="subsection-header">⚡ Maintenance Priority Analysis</div>',
            unsafe_allow_html=True,
        )
        fig_priority = create_maintenance_priority_chart(filtered_df)
        st.plotly_chart(fig_priority, use_container_width=True)

        # Priority explanation
        st.info(
            """
        **Priority Level Guide:**
        - 🔴 **90-100**: Overdue maintenance (immediate action required)
        - 🟠 **70-89**: Due within 30 days (schedule soon) 
        - 🟡 **50-69**: Due within 90 days (plan ahead)
        - 🟢 **30-49**: Long-term planning (3+ months)
        """
        )

    with tab2:
        st.markdown(
            '<div class="section-header">🔮 Predictive Analytics - Maintenance Demand Forecasting</div>',
            unsafe_allow_html=True,
        )

        # Load model for display info
        model, label_encoder, model_info = load_model()

        # Model performance info
        if model_info:
            colors = get_theme_colors()
            col1, col2, col3 = st.columns(3)
            with col1:
                st.markdown(
                    f"""
                    <div style="background: {colors['card_bg']}; 
                                padding: 1.5rem; 
                                border-radius: 12px; 
                                box-shadow: 0 4px 12px {colors['shadow']};
                                border: 1px solid {colors['border']};
                                text-align: center;
                                transition: all 0.3s ease;">
                        <p style="color: {colors['text_secondary']} !important; font-size: 0.875rem; margin: 0; font-weight: 500;">Model Type</p>
                        <h2 style="color: {colors['text_primary']} !important; margin: 0.5rem 0 0 0; font-weight: bold; font-size: 1.5rem;">{model_info["model_name"]}</h2>
                    </div>
                    """,
                    unsafe_allow_html=True,
                )
            with col2:
                st.markdown(
                    f"""
                    <div style="background: {colors['card_bg']}; 
                                padding: 1.5rem; 
                                border-radius: 12px; 
                                box-shadow: 0 4px 12px {colors['shadow']};
                                border: 1px solid {colors['border']};
                                text-align: center;
                                transition: all 0.3s ease;">
                        <p style="color: {colors['text_secondary']} !important; font-size: 0.875rem; margin: 0; font-weight: 500;">Accuracy</p>
                        <h2 style="color: {colors['text_primary']} !important; margin: 0.5rem 0 0 0; font-weight: bold; font-size: 1.5rem;">{model_info['metrics']['accuracy']:.2%}</h2>
                    </div>
                    """,
                    unsafe_allow_html=True,
                )
            with col3:
                st.markdown(
                    f"""
                    <div style="background: {colors['card_bg']}; 
                                padding: 1.5rem; 
                                border-radius: 12px; 
                                box-shadow: 0 4px 12px {colors['shadow']};
                                border: 1px solid {colors['border']};
                                text-align: center;
                                transition: all 0.3s ease;">
                        <p style="color: {colors['text_secondary']} !important; font-size: 0.875rem; margin: 0; font-weight: 500;">F1-Score</p>
                        <h2 style="color: {colors['text_primary']} !important; margin: 0.5rem 0 0 0; font-weight: bold; font-size: 1.5rem;">{model_info['metrics']['f1_score']:.3f}</h2>
                    </div>
                    """,
                    unsafe_allow_html=True,
                )

        st.markdown("---")

        # Forecast period selector
        col1, col2 = st.columns([1, 3])
        with col1:
            colors = get_theme_colors()
            st.markdown(
                f"""
                <div style="background: {colors['card_bg']}; 
                            padding: 1.2rem; 
                            border-radius: 12px; 
                            box-shadow: 0 4px 12px {colors['shadow']};
                            border: 1px solid {colors['border']};
                            transition: all 0.3s ease;">
                    <p style="color: {colors['text_primary']} !important; font-size: 0.875rem; margin: 0 0 0.8rem 0; font-weight: 500;">Forecast Period</p>
                """,
                unsafe_allow_html=True,
            )
            forecast_months = st.selectbox(
                "Forecast Period (months)",
                [3, 6, 9, 12],
                index=1,
                help="Select how many months ahead to forecast",
                key="forecast_period_select",
                label_visibility="collapsed",
            )
            st.markdown("</div>", unsafe_allow_html=True)

        st.markdown("### 📈 Maintenance Demand Forecast")

        # Generate forecast charts
        fig1, fig2, fig3, forecast_df = create_forecast_charts(
            filtered_df, forecast_months
        )

        # Display main forecast charts
        col1, col2 = st.columns(2)
        with col1:
            st.plotly_chart(fig1, use_container_width=True)
        with col2:
            st.plotly_chart(fig2, use_container_width=True)

        # Risk breakdown chart
        st.plotly_chart(fig3, use_container_width=True)

        st.markdown("### 🔍 Category-Specific Forecasts")

        # Category tabs for detailed analysis
        category_tabs = st.tabs(list(filtered_df["category"].unique()))

        for i, category in enumerate(filtered_df["category"].unique()):
            with category_tabs[i]:
                fig_cat, types_df = create_asset_type_forecast(
                    filtered_df, category, forecast_months
                )
                st.plotly_chart(fig_cat, use_container_width=True)

                # Top asset types needing attention
                top_types = (
                    types_df.groupby("Type")["Total_Risk"]
                    .sum()
                    .sort_values(ascending=False)
                    .head(5)
                )

                st.markdown(f"**Top 5 {category} Types Requiring Attention:**")
                for idx, (asset_type, risk_score) in enumerate(top_types.items(), 1):
                    st.markdown(
                        f"{idx}. **{asset_type}**: {int(risk_score)} assets at risk"
                    )

        st.markdown("---")

        # Generate insights and recommendations
        st.markdown("### 💡 Insights & Recommendations")

        insights, recommendations = generate_insights_and_recommendations(
            filtered_df, forecast_df
        )

        col1, col2 = st.columns(2)

        with col1:
            # Key Insights Card
            colors = get_theme_colors()
            insights_html = f"""
            <div style="background: {colors['card_bg']}; 
                        padding: 1.5rem; 
                        border-radius: 12px; 
                        box-shadow: 0 4px 12px {colors['shadow']};
                        border: 1px solid {colors['border']};
                        margin-bottom: 1rem;
                        transition: all 0.3s ease;">
                <h4 style="color: {colors['text_primary']} !important; margin-top: 0; margin-bottom: 1rem;">📊 Key Insights</h4>
            """

            for insight in insights:
                insights_html += f'<p style="color: {colors["text_primary"]} !important; margin-bottom: 0.5rem;">{insight}</p>'

            insights_html += "</div>"
            st.markdown(insights_html, unsafe_allow_html=True)

            # Statistical Summary Card
            total_assets_at_risk = forecast_df["Total_Demand"].sum()
            avg_monthly_demand = (
                forecast_df.groupby("Month")["Total_Demand"].sum().mean()
            )

            stats_html = f"""
            <div style="background: {colors['card_bg']}; 
                        padding: 1.5rem; 
                        border-radius: 12px; 
                        box-shadow: 0 4px 12px {colors['shadow']};
                        border: 1px solid {colors['border']};
                        margin-bottom: 1rem;
                        transition: all 0.3s ease;">
                <h4 style="color: {colors['text_primary']} !important; margin-top: 0; margin-bottom: 1rem;">📈 Statistical Summary</h4>
                <p style="color: {colors['text_primary']} !important; margin-bottom: 0.5rem;">• <strong>Total Assets at Risk</strong>: {int(total_assets_at_risk)} over {forecast_months} months</p>
                <p style="color: {colors['text_primary']} !important; margin-bottom: 0.5rem;">• <strong>Average Monthly Demand</strong>: {int(avg_monthly_demand)} assets</p>
                <p style="color: {colors['text_primary']} !important; margin-bottom: 0.5rem;">• <strong>Peak vs Average Ratio</strong>: {forecast_df.groupby("Month")["Total_Demand"].sum().max() / avg_monthly_demand:.1f}x</p>
            </div>
            """
            st.markdown(stats_html, unsafe_allow_html=True)

        with col2:
            # Action Recommendations Card
            recommendations_html = f"""
            <div style="background: {colors['card_bg']}; 
                        padding: 1.5rem; 
                        border-radius: 12px; 
                        box-shadow: 0 4px 12px {colors['shadow']};
                        border: 1px solid {colors['border']};
                        margin-bottom: 1rem;
                        transition: all 0.3s ease;">
                <h4 style="color: {colors['text_primary']} !important; margin-top: 0; margin-bottom: 1rem;">🎯 Action Recommendations</h4>
            """

            for recommendation in recommendations:
                recommendations_html += f'<p style="color: {colors["text_primary"]} !important; margin-bottom: 0.5rem;">{recommendation}</p>'

            recommendations_html += "</div>"
            st.markdown(recommendations_html, unsafe_allow_html=True)

            # Priority Actions Card
            priority_items = [
                "📅 Schedule maintenance teams for peak periods",
                "📦 Order spare parts for high-risk asset types",
                "👥 Train additional maintenance personnel",
                "💰 Allocate emergency maintenance budget",
                "🔄 Implement preventive maintenance cycles",
            ]

            priority_html = f"""
            <div style="background: {colors['card_bg']}; 
                        padding: 1.5rem; 
                        border-radius: 12px; 
                        box-shadow: 0 4px 12px {colors['shadow']};
                        border: 1px solid {colors['border']};
                        margin-bottom: 1rem;
                        transition: all 0.3s ease;">
                <h4 style="color: {colors['text_primary']} !important; margin-top: 0; margin-bottom: 1rem;">⚡ Priority Actions</h4>
            """

            for item in priority_items:
                priority_html += f'<p style="color: {colors["text_primary"]} !important; margin-bottom: 0.5rem;">• {item}</p>'

            priority_html += "</div>"
            st.markdown(priority_html, unsafe_allow_html=True)

        # Forecast summary table
        st.markdown("### 📋 Forecast Summary Table")

        # Debug: Check if forecast_df has data
        if forecast_df.empty:
            st.warning(
                "No forecast data available. This might be due to insufficient historical data."
            )
            st.info(
                "Please ensure your asset data has valid purchase dates and maintenance schedules."
            )
        else:
            # Show basic info about the data
            st.write(
                f"**Data Available**: {len(forecast_df)} records across {forecast_df['Month'].nunique()} months"
            )

            try:
                # Create summary table
                if "Month" in forecast_df.columns and "Category" in forecast_df.columns:
                    summary_table = (
                        forecast_df.groupby(["Month", "Category"])
                        .agg(
                            {
                                "Scheduled_Maintenance": "sum",
                                "High_Usage_Risk": "sum",
                                "Aging_Assets_Risk": "sum",
                                "Total_Demand": "sum",
                            }
                        )
                        .reset_index()
                    )

                    if len(summary_table) > 0:
                        st.write("**Forecast Summary:**")

                        # Use HTML table with inline styles that can't be overridden
                        html_table = "<div style='background: white; padding: 20px; border-radius: 8px; border: 1px solid #ddd;'>"
                        html_table += "<table style='width: 100%; border-collapse: collapse; font-family: Arial, sans-serif;'>"

                        # Header
                        html_table += "<thead><tr style='background-color: #f8f9fa;'>"
                        for col in summary_table.columns:
                            html_table += f"<th style='padding: 12px; border: 1px solid #ddd; color: #333; font-weight: bold; text-align: left;'>{col}</th>"
                        html_table += "</tr></thead>"

                        # Body
                        html_table += "<tbody>"
                        for idx, row in summary_table.iterrows():
                            bg_color = "#f8f9fa" if idx % 2 == 0 else "white"
                            html_table += f"<tr style='background-color: {bg_color};'>"
                            for value in row:
                                html_table += f"<td style='padding: 12px; border: 1px solid #ddd; color: #333;'>{value}</td>"
                            html_table += "</tr>"
                        html_table += "</tbody></table></div>"

                        st.markdown(html_table, unsafe_allow_html=True)

                    else:
                        st.error("Summary table is empty after grouping.")

                        # Fallback: show raw data as HTML table
                        st.write("**Raw Forecast Data (fallback):**")
                        html_table = "<div style='background: white; padding: 20px; border-radius: 8px; border: 1px solid #ddd;'>"
                        html_table += "<table style='width: 100%; border-collapse: collapse; font-family: Arial, sans-serif;'>"

                        # Header
                        html_table += "<thead><tr style='background-color: #f8f9fa;'>"
                        for col in forecast_df.columns:
                            html_table += f"<th style='padding: 12px; border: 1px solid #ddd; color: #333; font-weight: bold; text-align: left;'>{col}</th>"
                        html_table += "</tr></thead>"

                        # Body (show first 10 rows)
                        html_table += "<tbody>"
                        for idx, row in forecast_df.head(10).iterrows():
                            bg_color = "#f8f9fa" if idx % 2 == 0 else "white"
                            html_table += f"<tr style='background-color: {bg_color};'>"
                            for value in row:
                                html_table += f"<td style='padding: 12px; border: 1px solid #ddd; color: #333;'>{value}</td>"
                            html_table += "</tr>"
                        html_table += "</tbody></table></div>"

                        st.markdown(html_table, unsafe_allow_html=True)

                else:
                    st.error(
                        "Required columns 'Month' and/or 'Category' not found in forecast data."
                    )
                    st.write("Available columns:", forecast_df.columns.tolist())

            except Exception as e:
                st.error(f"Error creating forecast table: {str(e)}")
                st.write("**Error Details:**")
                st.write("- Exception type:", type(e).__name__)
                st.write("- Exception message:", str(e))

        # Export forecast data
        if st.button("📥 Export Forecast Data"):
            csv_data = forecast_df.to_csv(index=False)
            st.download_button(
                label="Download Forecast CSV",
                data=csv_data,
                file_name=f"maintenance_forecast_{forecast_months}months.csv",
                mime="text/csv",
            )

    with tab3:
        st.markdown(
            '<div class="section-header">🤖 AI Report Generator</div>',
            unsafe_allow_html=True,
        )

        st.markdown(
            """
        Generate comprehensive AI-powered insights and recommendations based on your asset data.
        """
        )

        # Report configuration
        st.markdown(
            '<div class="subsection-header">⚙️ Report Configuration</div>',
            unsafe_allow_html=True,
        )

        # Configuration options outside form
        col1, col2 = st.columns(2)

        with col1:
            include_kpis = st.checkbox("Include KPIs", value=True)
            include_trends = st.checkbox("Include Trend Analysis", value=True)
            include_distribution = st.checkbox(
                "Include Distribution Charts", value=True
            )

        with col2:
            include_heatmap = st.checkbox("Include Geographic Heatmap", value=True)
            include_forecast = st.checkbox("Include Maintenance Forecast", value=True)
            report_format = st.selectbox(
                "Report Format", ["Word Document", "PDF", "Text"], index=0
            )

        report_focus = st.selectbox(
            "Report Focus",
            [
                "Executive Summary",
                "Operational Analysis",
                "Strategic Planning",
                "Risk Assessment",
            ],
        )

        st.markdown("---")

        # Generate button - moved outside of form to fix the error
        if st.button("🚀 Generate AI Report", type="primary", use_container_width=True):
            report_config = {
                "include_kpis": include_kpis,
                "include_trends": include_trends,
                "include_distribution": include_distribution,
                "include_heatmap": include_heatmap,
                "include_forecast": include_forecast,
                "focus": report_focus,
            }

            with st.spinner("Generating AI-powered report... This may take a moment."):
                # Prepare charts data for Word export
                charts_data = {
                    "total_assets": len(filtered_df),
                    "good_assets": len(filtered_df[filtered_df["condition"] == "Good"]),
                    "needs_action": len(
                        filtered_df[filtered_df["condition"] == "Needs Action"]
                    ),
                    "damaged_assets": len(
                        filtered_df[filtered_df["condition"] == "Damaged"]
                    ),
                    "overdue_maintenance": len(
                        filtered_df[
                            pd.to_datetime(filtered_df["next_maintenance_due"])
                            < datetime.now()
                        ]
                    ),
                    "good_pct": (
                        len(filtered_df[filtered_df["condition"] == "Good"])
                        / len(filtered_df)
                    )
                    * 100,
                    "needs_action_pct": (
                        len(filtered_df[filtered_df["condition"] == "Needs Action"])
                        / len(filtered_df)
                    )
                    * 100,
                    "damaged_pct": (
                        len(filtered_df[filtered_df["condition"] == "Damaged"])
                        / len(filtered_df)
                    )
                    * 100,
                    "category_stats": {},
                }

                # Add category statistics
                for category in filtered_df["category"].unique():
                    cat_data = filtered_df[filtered_df["category"] == category].copy()
                    cat_data["asset_age_years"] = (
                        pd.to_datetime("today")
                        - pd.to_datetime(cat_data["purchase_date"])
                    ).dt.days / 365.25

                    charts_data["category_stats"][category] = {
                        "total_assets": len(cat_data),
                        "good_count": len(cat_data[cat_data["condition"] == "Good"]),
                        "needs_action_count": len(
                            cat_data[cat_data["condition"] == "Needs Action"]
                        ),
                        "damaged_count": len(
                            cat_data[cat_data["condition"] == "Damaged"]
                        ),
                        "good_pct": (
                            len(cat_data[cat_data["condition"] == "Good"])
                            / len(cat_data)
                        )
                        * 100,
                        "needs_action_pct": (
                            len(cat_data[cat_data["condition"] == "Needs Action"])
                            / len(cat_data)
                        )
                        * 100,
                        "damaged_pct": (
                            len(cat_data[cat_data["condition"] == "Damaged"])
                            / len(cat_data)
                        )
                        * 100,
                        "avg_age_years": cat_data["asset_age_years"].mean(),
                        "top_types": cat_data["type"].value_counts(),
                    }

                # Generate the report
                report_content = generate_ai_report(report_config, kpis, filtered_df)

                if report_content:
                    st.success("✅ Report generated successfully!")

                    # Display the report
                    st.markdown(
                        '<div class="subsection-header">📄 Generated Report</div>',
                        unsafe_allow_html=True,
                    )
                    st.markdown(report_content)

                    st.markdown("---")
                    st.markdown(
                        '<div class="subsection-header">📥 Download Report</div>',
                        unsafe_allow_html=True,
                    )

                    # Download buttons
                    col1, col2, col3 = st.columns(3)

                    with col1:
                        if report_format == "Word Document":
                            word_content = export_report_to_word_with_charts(
                                report_content, charts_data, df
                            )
                            if word_content:
                                st.download_button(
                                    label="📄 Download Word Report",
                                    data=word_content,
                                    file_name=f"PDRM_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    type="primary",
                                )
                        else:
                            st.info(
                                "Select Word Document format to download with charts"
                            )

                    with col2:
                        if report_format == "PDF":
                            pdf_content = export_report_to_pdf(report_content)
                            if pdf_content:
                                st.download_button(
                                    label="📄 Download PDF Report",
                                    data=pdf_content,
                                    file_name=f"PDRM_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
                                    mime="application/pdf",
                                    type="primary",
                                )
                        else:
                            st.info("Select PDF format to download PDF")

                    with col3:
                        st.download_button(
                            label="📝 Download Text Report",
                            data=report_content,
                            file_name=f"PDRM_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                            mime="text/plain",
                        )

                else:
                    st.error(
                        "Failed to generate report. Please check your OpenAI API key and try again."
                    )

        # Expired Stock List Report Section
        st.markdown("---")
        st.markdown(
            '<div class="subsection-header">📋 Laporan Senarai Stok Bertarikh Luput (KEW.PS-6)</div>',
            unsafe_allow_html=True,
        )

        st.markdown(
            """
        Jana **Senarai Stok Bertarikh Luput** mengikut borang rasmi kerajaan Malaysia KEW.PS-6. 
        Laporan ini merangkumi aset yang rosak atau mempunyai jadual penyelenggaraan yang tertunggak.
        """
        )

        col1, col2, col3 = st.columns([1, 1, 2])

        with col1:
            if st.button(
                "📄 Jana Laporan Stok Bertarikh Luput (KEW.PS-6)",
                type="secondary",
                use_container_width=True,
            ):
                with st.spinner("Menghasilkan laporan Senarai Stok Bertarikh Luput..."):
                    expired_report = generate_expired_stock_report(filtered_df)

                    if expired_report:
                        st.success(
                            "✅ Laporan Stok Bertarikh Luput berjaya dihasilkan!"
                        )

                        st.download_button(
                            label="📄 Muat Turun Senarai Stok Bertarikh Luput",
                            data=expired_report,
                            file_name=f"Senarai_Stok_Bertarikh_Luput_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            type="primary",
                        )
                    else:
                        st.error("Gagal menghasilkan laporan stok bertarikh luput.")

        with col2:
            # Show summary of items that would be in the report
            expired_items = filtered_df[
                (filtered_df["condition"] == "Damaged")
                | (
                    pd.to_datetime(filtered_df["next_maintenance_due"], errors="coerce")
                    < datetime.now()
                )
            ]

            st.info(
                f"""
            **Report Preview:**
            - Damaged Assets: {len(expired_items[expired_items['condition'] == 'Damaged'])}
            - Overdue Maintenance: {len(expired_items[pd.to_datetime(expired_items['next_maintenance_due'], errors='coerce') < datetime.now()])}
            - Total Items: {len(expired_items)}
            """
            )

        # Sample report preview
        st.markdown("---")
        st.markdown(
            '<div class="subsection-header">📋 Sample Report Preview</div>',
            unsafe_allow_html=True,
        )

        with st.expander("View Sample Report Structure"):
            st.markdown(
                """
            **EXECUTIVE SUMMARY**
            - Asset overview and key metrics
            - Critical issues identification
            - High-priority recommendations
            
            **KEY FINDINGS AND INSIGHTS**
            - Category-wise condition analysis
            - Geographic distribution patterns
            - Age and maintenance correlations
            
            **ASSET CONDITION ANALYSIS**
            - Detailed condition breakdown by category
            - Risk factors and patterns
            - Performance indicators
            
            **MAINTENANCE RECOMMENDATIONS**
            - Overdue maintenance priorities
            - Predictive maintenance strategies
            - Resource allocation suggestions
            
            **RISK ASSESSMENT**
            - Critical asset identification
            - Failure probability analysis
            - Operational impact assessment
            
            **STRATEGIC RECOMMENDATIONS**
            - Investment priorities
            - Policy improvements
            - Long-term planning insights
            """
            )

        # AI Configuration
        with st.expander("🔧 Advanced Settings"):
            st.markdown("**AI Report Configuration**")

            col1, col2 = st.columns(2)
            with col1:
                st.info("**Report Focus Areas:**")
                st.write("• Executive Summary: High-level overview for leadership")
                st.write("• Operational Analysis: Day-to-day operational insights")
                st.write("• Strategic Planning: Long-term strategic recommendations")
                st.write("• Risk Assessment: Risk-focused analysis and mitigation")

            with col2:
                st.info("**Export Options:**")
                st.write("• Word Document: Full report with charts and formatting")
                st.write("• PDF: Professional document format")
                st.write("• Text: Plain text for further editing")

            st.warning(
                "**Note:** AI report generation requires a valid OpenAI API key. Set your API key as an environment variable 'OPENAI_API_KEY'."
            )

    # AI Assistant Chatbot
    render_ai_chatbot(filtered_df)

    # Footer
    st.markdown("---")
    st.markdown(
        """
        <div style='text-align: center; color: #666;'>
            <p>Jabatan Logistik & Asset - Powered By Credence Analytics & AI | Royal Malaysia Police</p>
            <p>For support, contact: asset.management@pdrm.gov.my</p>
        </div>
        """,
        unsafe_allow_html=True,
    )


if __name__ == "__main__":
    main()
